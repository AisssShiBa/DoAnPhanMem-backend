from pathlib import Path
from shutil import copy2

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt


TARGET = Path(r"D:\GK_AI\CuoiKiAI.docx")
BACKUP = TARGET.with_name("CuoiKiAI_backup_truoc_khi_chuan_hoa_cong_thuc_lan2.docx")
OUTPUT_IF_LOCKED = TARGET.with_name("CuoiKiAI_DA_CHUAN_HOA_CONG_THUC_LAN2.docx")


REPL = {
    "J(θ)=1/(2m) sum(h_θ(xᵢ)-yᵢ)^2": "J(θ) = (1/(2m))∑ᵢ(h_θ(xᵢ) − yᵢ)²",
    "J(θ)=1/(2m) ∑ᵢ (h_θ(xᵢ)-yᵢ)^2": "J(θ) = (1/(2m))∑ᵢ(h_θ(xᵢ) − yᵢ)²",
    "J(θ)=1/(2m) ∑ᵢ (h_θ(xᵢ)-yᵢ)²": "J(θ) = (1/(2m))∑ᵢ(h_θ(xᵢ) − yᵢ)²",
    "J(θ)=1/(2m) ∑ᵢ(h_θ(xᵢ)-yᵢ)²": "J(θ) = (1/(2m))∑ᵢ(h_θ(xᵢ) − yᵢ)²",
    "g(z)=1/(1+e^(-z))": "g(z) = 1/(1 + e⁻ᶻ)",
    "Q:=Q+α[r+γ max Q(s',a')-Q]": "Q(s,a) ← Q(s,a) + α[r + γ maxₐ′ Q(s′,a′) − Q(s,a)]",
    "Q(s,a):=Q(s,a)+α[r+γ max_a' Q(s',a')-Q(s,a)]": "Q(s,a) ← Q(s,a) + α[r + γ maxₐ′ Q(s′,a′) − Q(s,a)]",
    "Q(s,a):=Q(s,a)+α[r+γ maxₐ' Q(s',a')-Q(s,a)]": "Q(s,a) ← Q(s,a) + α[r + γ maxₐ′ Q(s′,a′) − Q(s,a)]",
    "r + γ max_a' Q(s',a')": "r + γ maxₐ′ Q(s′,a′)",
    "max_a'": "maxₐ′",
    "s'": "s′",
    "a'": "a′",
    "PositiveSet(hⱼ) subset PositiveSet(hᵢ)": "PositiveSet(hⱼ) ⊆ PositiveSet(hᵢ)",
    "subset": "⊆",
    "H(S) = - ∑ᵢ pᵢ log₂(pᵢ)": "H(S) = −∑ᵢ pᵢlog₂(pᵢ)",
    "H(S) = −∑ᵢ pᵢ log₂(pᵢ)": "H(S) = −∑ᵢ pᵢlog₂(pᵢ)",
    "−∑ᵢ pᵢ log₂(pᵢ)": "−∑ᵢ pᵢlog₂(pᵢ)",
    "H(S)=-∑ᵢ pᵢ log₂(pᵢ)": "H(S) = −∑ᵢ pᵢlog₂(pᵢ)",
    "Gini(S)=1-": "Gini(S) = 1 − ",
    "1-p^2-(1-p)^2": "1 − p² − (1−p)²",
    "0.6^2": "0.6²",
    "0.4^2": "0.4²",
    "Bias(f̂(x₀))]^2": "Bias(f̂(x₀))]²",
    "[Bias(f̂(x₀))]^2": "Bias(f̂(x₀))²",
    "Bias(f̂)^2": "Bias(f̂)²",
    "f̂(x₀))^2": "f̂(x₀))²",
    "Y₀ -": "Y₀ −",
    "1/(ε)": "(1/ε)",
    "(1/ε)(ln|H|+ln(1/δ))": "(1/ε)(ln|H| + ln(1/δ))",
    "pᵢ log₂": "pᵢlog₂",
    "p₊log₂": "p₊log₂",
    "p₋log₂": "p₋log₂",
}


def normalize_font(doc):
    for style_name in ["Normal", "Heading 1", "Heading 2", "Heading 3"]:
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        if style_name == "Normal":
            style.font.size = Pt(12)


def transform(text):
    new = text
    for old, repl in REPL.items():
        new = new.replace(old, repl)
    if "MSE" in new and "sum(" in new:
        new = new.replace("sum(", "∑ᵢ(").replace(")^2", ")²")
    return new


def set_text(paragraph, text):
    if paragraph.text == text:
        return 0
    if not paragraph.runs:
        paragraph.add_run(text)
    else:
        for run in paragraph.runs:
            run.text = ""
        paragraph.runs[0].text = text
    return 1


def main():
    if not TARGET.exists():
        raise FileNotFoundError(TARGET)
    if not BACKUP.exists():
        copy2(TARGET, BACKUP)
    doc = Document(TARGET)
    normalize_font(doc)
    changed = 0
    for para in doc.paragraphs:
        changed += set_text(para, transform(para.text))
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    changed += set_text(para, transform(para.text))
    try:
        doc.save(TARGET)
        print(f"UPDATED={TARGET}")
    except PermissionError:
        doc.save(OUTPUT_IF_LOCKED)
        print(f"TARGET_LOCKED={TARGET}")
        print(f"OUTPUT={OUTPUT_IF_LOCKED}")
    print(f"BACKUP={BACKUP}")
    print(f"CHANGED={changed}")


if __name__ == "__main__":
    main()
