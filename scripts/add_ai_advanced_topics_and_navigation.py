from pathlib import Path
from shutil import copy2

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


TARGET = Path(r"D:\GK_AI\CuoiKiAI.docx")
BACKUP = TARGET.with_name("CuoiKiAI_backup_truoc_khi_toi_uu_tra_cuu.docx")
OUTPUT_IF_LOCKED = TARGET.with_name("CuoiKiAI_DA_BO_SUNG_DAY_DU_DE_MO.docx")


def normalize_font(doc: Document):
    for style_name in ["Normal", "Heading 1", "Heading 2", "Heading 3"]:
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        if style_name == "Normal":
            style.font.size = Pt(12)


def p(doc, text="", label=None, indent=True):
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(4)
    para.paragraph_format.line_spacing = 1.08
    if indent:
        para.paragraph_format.first_line_indent = Inches(0.25)
    if label and text.startswith(label):
        run = para.add_run(label)
        run.bold = True
        para.add_run(text[len(label):])
    else:
        para.add_run(text)
    return para


def h1(doc, text):
    para = doc.add_paragraph(style="Heading 1")
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    return para


def h2(doc, text):
    para = doc.add_paragraph(style="Heading 2")
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(13)
    return para


def h3(doc, text):
    para = doc.add_paragraph(style="Heading 3")
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    return para


def bullet(doc, text):
    para = doc.add_paragraph(style="List Bullet")
    para.paragraph_format.space_after = Pt(2)
    para.add_run(text)
    return para


def table(doc, headers, rows):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, header in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        cell.text = header
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for row in rows:
        cells = tbl.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    doc.add_paragraph()
    return tbl


def style_existing_headings(doc: Document):
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        if text.startswith("PHẦN "):
            para.style = "Heading 1"
        elif text.startswith("Câu mở rộng") or text.startswith("Câu "):
            para.style = "Heading 2"
        elif (
            text[:2].isdigit()
            and ". " in text[:5]
            and len(text) < 120
        ):
            para.style = "Heading 2"
        elif (
            len(text) < 90
            and (
                text.startswith("Chương ")
                or text.startswith("TÀI LIỆU THAM KHẢO")
                or text.startswith("MỤC ")
            )
        ):
            para.style = "Heading 2"


def prepend_navigation_note(doc: Document):
    if any("HƯỚNG DẪN TRA CỨU NHANH TRONG PHÒNG THI" in x.text for x in doc.paragraphs):
        return

    first = doc.paragraphs[0]._p
    body = first.getparent()
    insert_items = [
        ("Heading 1", "HƯỚNG DẪN TRA CỨU NHANH TRONG PHÒNG THI"),
        ("Normal", "File này đã được tối ưu cho đề mở: dùng Ctrl+F theo các mã/từ khóa như [CH2-BIAS], [CH3-PROLOG-EBG], [CH5-CI], [CH5-LOOCV], [CH5-BOOTSTRAP], A*, entropy, overfitting, Q-learning."),
        ("Normal", "Cách dùng nhanh: đọc đề -> tìm từ khóa -> mở đúng phần -> chép theo khung Định nghĩa + Công thức/thuật toán + Ví dụ + So sánh + Kết luận."),
        ("Normal", "Trong Word có thể mở Navigation Pane để nhảy theo Heading. Các mục PHẦN, Câu, Câu mở rộng đã được gắn style Heading để tìm nhanh hơn."),
    ]
    for style, text in reversed(insert_items):
        new_p = OxmlElement("w:p")
        first.addprevious(new_p)
        para = doc.paragraphs[0]
        para.style = style
        run = para.add_run(text)
        if style.startswith("Heading"):
            run.bold = True
        para.paragraph_format.space_after = Pt(4)


def already_added(doc: Document) -> bool:
    return any("PHẦN 6. BỔ SUNG CHUYÊN SÂU THEO MITCHELL VÀ ISLR" in p.text for p in doc.paragraphs)


def add_advanced_section(doc: Document):
    doc.add_page_break()
    h1(doc, "PHẦN 6. BỔ SUNG CHUYÊN SÂU THEO MITCHELL VÀ ISLR")
    p(doc, "Phần này bổ sung các ý nâng cao rất đáng có cho đề mở. Các mục được gắn mã để Ctrl+F nhanh: [CH2-BIAS], [CH2-QUERY], [CH3-PROLOG-EBG], [CH3-FOIL-GAIN], [CH5-ERROR], [CH5-CI], [CH5-LOOCV], [CH5-BOOTSTRAP].")

    h2(doc, "1. [CH2-BIAS] Inductive Bias - Độ chệch quy nạp")
    p(doc, "Định nghĩa: Inductive bias là tập hợp các giả định mà thuật toán học dùng để dự đoán nhãn cho những ví dụ chưa từng thấy. Nếu không có inductive bias, từ một tập huấn luyện hữu hạn sẽ có vô số giả thuyết cùng khớp dữ liệu nhưng dự đoán khác nhau trên dữ liệu mới.", "Định nghĩa:")
    p(doc, "Ý nghĩa: Học máy không chỉ là ghi nhớ dữ liệu. Muốn tổng quát hóa, thuật toán phải có giả định nào đó về dạng của khái niệm mục tiêu, cách ưu tiên giả thuyết hoặc phạm vi không gian giả thuyết.")
    p(doc, "Ví dụ Candidate-Elimination: Bias của Candidate-Elimination là khái niệm mục tiêu nằm trong hypothesis space H và dữ liệu huấn luyện không nhiễu. Nếu target concept không nằm trong H hoặc dữ liệu bị gán nhãn sai, version space có thể rỗng hoặc không chứa giả thuyết đúng.", "Ví dụ Candidate-Elimination:")
    p(doc, "Ví dụ dễ nhớ: Nếu H chỉ gồm các hình chữ nhật trên mặt phẳng, thuật toán chỉ có thể học ranh giới dạng chữ nhật. Khi khái niệm thật là hình tròn, dù dữ liệu nhiều hơn thì thuật toán vẫn bị giới hạn bởi bias của H.")
    p(doc, "Cách đưa vào bài thi: Khi đề hỏi vì sao thuật toán học có thể dự đoán ví dụ chưa thấy, hãy trả lời rằng nó dựa vào inductive bias. Bias không hẳn xấu; nó là điều kiện cần để tổng quát hóa, nhưng bias sai sẽ gây underfitting hoặc học sai quy luật.")

    h2(doc, "2. [CH2-QUERY] Optimal Query Strategy trong học khái niệm")
    p(doc, "Khái niệm: Optimal Query Strategy là chiến lược chọn ví dụ truy vấn sao cho thu hẹp version space nhanh nhất. Ý tưởng tốt nhất là chọn ví dụ mà các giả thuyết còn lại trong version space bất đồng mạnh nhất, lý tưởng là chia version space thành hai nửa gần bằng nhau.")
    p(doc, "Liên hệ tìm kiếm nhị phân: Nếu mỗi truy vấn loại được khoảng một nửa giả thuyết còn lại, số truy vấn cần thiết giảm rất nhanh, tương tự binary search. Đây là cách học chủ động hơn so với chờ ví dụ ngẫu nhiên.")
    p(doc, "Ví dụ cụ thể: Giả sử version space còn 8 giả thuyết. Nếu chọn một ví dụ mà 4 giả thuyết dự đoán dương và 4 giả thuyết dự đoán âm, sau khi biết nhãn thật ta loại được khoảng 4 giả thuyết. Nếu chọn ví dụ mà 7 giả thuyết cùng dự đoán dương và 1 giả thuyết dự đoán âm, thông tin thu được ít hơn nhiều.")
    p(doc, "Cách viết kết luận: Truy vấn tối ưu không nhất thiết là ví dụ dễ phân loại nhất; nó là ví dụ nhiều thông tin nhất vì làm các giả thuyết còn lại bộc lộ sự khác nhau.")

    h2(doc, "3. [CH3-PROLOG-EBG] PROLOG-EBG và Weakest Preimage")
    p(doc, "Khái niệm: PROLOG-EBG là một dạng cụ thể của Explanation-Based Learning. Nó dùng cơ chế suy diễn kiểu Prolog, thường là backward chaining, để xây dựng proof cho một ví dụ huấn luyện, sau đó tổng quát hóa proof đó thành một luật mới.")
    p(doc, "Quy trình: Bước 1 nhận ví dụ dương cần học. Bước 2 dùng domain theory để chứng minh ví dụ bằng backward chaining. Bước 3 phân tích proof tree để xác định các điều kiện nào thật sự cần thiết. Bước 4 tính weakest preimage, tức tập điều kiện yếu nhất nhưng vẫn đủ để đảm bảo kết luận. Bước 5 tạo rule tổng quát từ các điều kiện đó.")
    p(doc, "Weakest Preimage: Có thể hiểu là mô tả tổng quát nhất của đầu vào sao cho proof vẫn thành công. Nó không giữ lại chi tiết thừa của ví dụ cụ thể, mà chỉ giữ các điều kiện cần cho kết luận.")
    p(doc, "Ví dụ cụ thể: Nếu hệ biết Cup(x) and Light(x) -> Liftable(x), và biết Cup(obj1), Light(obj1), Color(obj1,red). Khi chứng minh Liftable(obj1), màu đỏ không nằm trong proof cần thiết. Weakest preimage chỉ giữ Cup(x) and Light(x), tạo luật tổng quát Cup(x) and Light(x) -> Liftable(x).")
    p(doc, "Ưu điểm và hạn chế: PROLOG-EBG học nhanh từ ít ví dụ vì tận dụng domain theory. Nhưng nếu domain theory sai, thiếu hoặc quá chi tiết, luật học được có thể sai, quá hẹp hoặc không hữu ích.")

    h2(doc, "4. [CH3-FOIL-GAIN] FOIL-Gain trong thuật toán FOIL")
    p(doc, "Khái niệm: FOIL học các luật logic dạng Horn clause bằng cách thêm dần literal vào phần điều kiện. FOIL-Gain là chỉ số dùng để chọn literal tốt nhất tại mỗi bước chuyên biệt hóa luật.")
    p(doc, "Trực giác: Literal tốt là literal làm luật bao phủ tỷ lệ ví dụ dương cao hơn và loại được nhiều ví dụ âm hơn. Nói cách khác, nó làm luật trở nên chính xác hơn nhưng vẫn giữ được nhiều ví dụ dương.")
    p(doc, "Công thức thường dùng: FOIL-Gain(L,R) = t * [log2(p1/(p1+n1)) - log2(p0/(p0+n0))]. Trong đó p0,n0 là số ví dụ dương/âm được luật cũ bao phủ; p1,n1 là số ví dụ dương/âm sau khi thêm literal L; t là số ví dụ dương vẫn còn được luật mới bao phủ.")
    p(doc, "Ví dụ cụ thể: Luật cũ bao phủ 6 ví dụ dương và 6 ví dụ âm, độ chính xác 6/12. Sau khi thêm literal A, luật bao phủ 5 dương và 1 âm, độ chính xác 5/6. Sau khi thêm literal B, luật bao phủ 3 dương và 0 âm, độ chính xác 3/3. FOIL-Gain cân nhắc cả độ sạch và số ví dụ dương còn giữ lại, nên không chỉ chọn literal có accuracy cao nhất nếu nó giữ lại quá ít ví dụ.")
    p(doc, "Cách kết luận: FOIL-Gain giống Information Gain trong cây quyết định ở chỗ đều chọn phép tách/điều kiện làm tăng thông tin, nhưng FOIL-Gain dùng trong học luật logic first-order.")

    h2(doc, "5. [CH5-ERROR] True Error và Sample Error trong đánh giá giả thuyết")
    p(doc, "True error error_D(h): là xác suất giả thuyết h phân loại sai trên toàn bộ phân phối dữ liệu D. Đây là sai số thật ta muốn biết, nhưng thường không đo trực tiếp được vì không có toàn bộ quần thể dữ liệu.")
    p(doc, "Sample error error_S(h): là tỷ lệ sai của h trên một tập mẫu S hữu hạn, ví dụ tập test. Đây là đại lượng đo được trong thực tế và dùng để ước lượng true error.")
    p(doc, "Công thức sample error cho phân loại: error_S(h) = số mẫu trong S bị h phân loại sai / tổng số mẫu trong S.")
    p(doc, "Ví dụ cụ thể: Một mô hình phân loại sai 12 email trong tập test 100 email. Khi đó error_S(h)=12/100=0.12. Nhưng true error trên toàn bộ email thực tế có thể không đúng bằng 0.12; ta cần khoảng tin cậy để ước lượng mức sai lệch có thể có.")
    p(doc, "Cách viết trong bài: Sample error là ước lượng của true error. Tập test càng lớn và càng đại diện cho phân phối D thì sample error càng đáng tin cậy.")

    h2(doc, "6. [CH5-CI] Confidence Interval cho sai số giả thuyết")
    p(doc, "Mục tiêu: Vì sample error chỉ được đo trên mẫu hữu hạn, ta dùng confidence interval để ước lượng khoảng giá trị có thể chứa true error với một mức tin cậy nhất định.")
    p(doc, "Công thức xấp xỉ: error_D(h) nằm trong khoảng error_S(h) ± z_N * sqrt(error_S(h)(1-error_S(h))/n). Trong đó n là số mẫu test, z_N phụ thuộc mức tin cậy: khoảng 1.64 cho 90%, 1.96 cho 95%, 2.58 cho 99%.")
    p(doc, "Ví dụ tính nhanh: Nếu error_S=0.12, n=100, mức tin cậy 95% dùng z=1.96. Độ lệch = 1.96 * sqrt(0.12*0.88/100) = 1.96 * sqrt(0.001056) ≈ 1.96 * 0.0325 ≈ 0.064. Vậy true error xấp xỉ nằm trong 0.12 ± 0.064, tức từ 0.056 đến 0.184.")
    p(doc, "Cách diễn giải: Không nên nói chắc chắn true error bằng sample error. Nên nói với mức tin cậy khoảng 95%, sai số thật được ước lượng nằm trong khoảng trên, nếu mẫu test độc lập và đại diện.")
    p(doc, "Lưu ý đi thi: Nếu n tăng, căn bậc hai ở mẫu số tăng làm khoảng tin cậy hẹp hơn. Nghĩa là test set lớn hơn cho ước lượng ổn định hơn.")

    h2(doc, "7. [CH5-LOOCV] LOOCV và so sánh với k-fold Cross-Validation")
    p(doc, "Khái niệm: Leave-One-Out Cross-Validation là trường hợp đặc biệt của k-fold cross-validation với k=n. Nếu có n mẫu, mỗi lần dùng n-1 mẫu để train và 1 mẫu để validation; lặp n lần rồi lấy trung bình lỗi.")
    p(doc, "Ưu điểm: LOOCV dùng gần như toàn bộ dữ liệu để huấn luyện ở mỗi lần, nên bias của ước lượng test error thường thấp, đặc biệt hữu ích khi dữ liệu ít.")
    p(doc, "Nhược điểm: Chi phí tính toán cao vì phải train mô hình n lần. Ngoài ra, các tập train ở mỗi lần rất giống nhau, nên các kết quả validation có thể tương quan mạnh và variance của ước lượng có thể lớn.")
    p(doc, "So sánh với k-fold: k-fold thường dùng k=5 hoặc k=10, train ít lần hơn nên rẻ hơn. Bias có thể cao hơn LOOCV một chút vì mỗi lần train dùng ít dữ liệu hơn, nhưng variance thường ổn định hơn trong nhiều tình huống thực tế.")
    p(doc, "Ví dụ cụ thể: Có 100 mẫu. LOOCV train 100 lần, mỗi lần 99 train và 1 validation. 10-fold CV train 10 lần, mỗi lần 90 train và 10 validation. Nếu mô hình rất tốn thời gian, 10-fold thường thực tế hơn.")

    h2(doc, "8. [CH5-BOOTSTRAP] Bootstrap trong ISLR")
    p(doc, "Khái niệm: Bootstrap là phương pháp lấy mẫu lại có thay thế từ tập dữ liệu gốc. Từ n quan sát ban đầu, ta tạo nhiều bootstrap samples, mỗi sample cũng có n quan sát nhưng một số điểm có thể xuất hiện nhiều lần và một số điểm không xuất hiện.")
    p(doc, "Mục đích: Bootstrap thường dùng để ước lượng độ không ổn định của một thống kê hoặc mô hình, ví dụ standard error của hệ số hồi quy, độ biến thiên của accuracy hoặc khoảng tin cậy cho một tham số.")
    p(doc, "Quy trình: Bước 1 lấy mẫu có thay thế từ dữ liệu gốc để tạo B tập bootstrap. Bước 2 huấn luyện hoặc tính thống kê trên mỗi tập. Bước 3 quan sát phân phối của B kết quả để ước lượng standard error hoặc confidence interval.")
    p(doc, "Ví dụ cụ thể: Có dữ liệu giá nhà và muốn biết hệ số của diện tích trong hồi quy tuyến tính ổn định không. Ta tạo 1000 bootstrap samples, fit mô hình 1000 lần, thu được 1000 hệ số diện tích. Nếu các hệ số dao động ít, ước lượng ổn định; nếu dao động mạnh, kết luận cần thận trọng.")
    p(doc, "So sánh với cross-validation: Cross-validation chủ yếu dùng để ước lượng test error và chọn mô hình. Bootstrap chủ yếu dùng để đo độ bất định/standard error của tham số hoặc thống kê, dù cũng có biến thể dùng đánh giá mô hình.")

    h2(doc, "9. Bảng so sánh nhanh các phần vừa bổ sung")
    table(doc, ["Khái niệm", "Dùng khi đề hỏi", "Câu chốt để ghi điểm"], [
        ("Inductive Bias", "Vì sao học được từ dữ liệu hữu hạn?", "Không có bias thì không thể tổng quát hóa từ mẫu hữu hạn."),
        ("Optimal Query Strategy", "Chọn ví dụ nào để hỏi nhãn?", "Chọn ví dụ chia version space gần 50-50 để thu hẹp nhanh."),
        ("PROLOG-EBG", "EBL hoạt động cụ thể thế nào?", "Dùng backward chaining tạo proof rồi lấy weakest preimage để tổng quát hóa."),
        ("FOIL-Gain", "FOIL chọn literal ra sao?", "Chọn literal tăng thông tin, giảm âm nhưng giữ được nhiều dương."),
        ("True vs Sample Error", "Đánh giá giả thuyết", "Sample error là ước lượng của true error trên mẫu hữu hạn."),
        ("Confidence Interval", "Bài tập tính sai số", "error_S ± z sqrt(error_S(1-error_S)/n)."),
        ("LOOCV", "Cross-validation nâng cao", "k=n, bias thấp nhưng tốn tính toán và có thể variance cao."),
        ("Bootstrap", "Ước lượng standard error", "Lấy mẫu lại có thay thế để đo độ bất định của thống kê/mô hình."),
    ])

    h2(doc, "10. Câu trả lời mẫu ghép nhiều ý nâng cao")
    p(doc, "Mẫu câu về đánh giá mô hình: Khi đánh giá một giả thuyết h, ta thật sự quan tâm true error error_D(h), tức sai số trên phân phối dữ liệu thật. Tuy nhiên ta chỉ đo được sample error error_S(h) trên tập test hữu hạn. Vì vậy cần dùng confidence interval để diễn tả độ bất định của ước lượng. Nếu sample error là 0.12 trên 100 mẫu, khoảng tin cậy 95% xấp xỉ là 0.12 ± 1.96*sqrt(0.12*0.88/100), tức khoảng 5.6% đến 18.4%.")
    p(doc, "Mẫu câu về học khái niệm: Trong concept learning, version space là tập giả thuyết còn nhất quán với dữ liệu. Candidate-Elimination thu hẹp version space bằng hai biên S và G, nhưng nó có inductive bias rằng target concept nằm trong H và dữ liệu không nhiễu. Nếu được chủ động truy vấn, nên chọn ví dụ làm các giả thuyết còn lại bất đồng nhiều nhất để loại được khoảng một nửa version space.")
    p(doc, "Mẫu câu về logic kết hợp học máy: EBL sử dụng tri thức nền để giải thích ví dụ, sau đó tổng quát hóa lời giải thích thành luật. PROLOG-EBG thực hiện ý tưởng này bằng backward chaining để tạo proof, rồi tính weakest preimage để loại bỏ chi tiết thừa và giữ lại điều kiện tổng quát nhất đủ cho kết luận. FOIL thì học luật logic theo hướng quy nạp, chọn literal bằng FOIL-Gain để tăng độ tinh khiết của luật.")


def main():
    if not TARGET.exists():
        raise FileNotFoundError(TARGET)
    if not BACKUP.exists():
        copy2(TARGET, BACKUP)

    doc = Document(TARGET)
    normalize_font(doc)
    style_existing_headings(doc)
    prepend_navigation_note(doc)
    if not already_added(doc):
        add_advanced_section(doc)
    try:
        doc.save(TARGET)
        print(f"UPDATED={TARGET}")
    except PermissionError:
        doc.save(OUTPUT_IF_LOCKED)
        print(f"TARGET_LOCKED={TARGET}")
        print(f"OUTPUT={OUTPUT_IF_LOCKED}")
    print(f"BACKUP={BACKUP}")
    print("ADDED_TOPICS=8")


if __name__ == "__main__":
    main()
