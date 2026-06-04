from pathlib import Path
from shutil import copy2

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Pt


TARGET = Path(r"D:\GK_AI\CuoiKiAI.docx")
BACKUP = TARGET.with_name("CuoiKiAI_backup_truoc_khi_chuan_hoa_van_phong.docx")
OUTPUT_IF_LOCKED = TARGET.with_name("CuoiKiAI_DA_CHUAN_HOA_VA_BO_SUNG.docx")


REPLACEMENTS = {
    "Câu kết luận nên chép:": "Kết luận chuẩn:",
    "Cách đưa vào bài thi:": "Cách trình bày trong bài thi:",
    "hãy trả lời rằng": "có thể trình bày rằng",
    "Câu chốt:": "Ý kết luận:",
    "Kernel không làm SVM 'ma thuật' hơn": "Kernel không làm SVM trở thành một phương pháp mơ hồ",
    "Mẫu câu về": "Đoạn trả lời mẫu về",
    "Mẫu câu trả lời khó đã ghép sẵn": "Các đoạn trả lời mẫu cho câu hỏi khó",
}


def normalize_font(doc: Document):
    for style_name in ["Normal", "Heading 1", "Heading 2", "Heading 3"]:
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        if style_name == "Normal":
            style.font.size = Pt(12)


def replace_text_preserve_runs(doc: Document):
    for para in doc.paragraphs:
        if not para.runs:
            continue
        original = para.text
        new = original
        for old, repl in REPLACEMENTS.items():
            new = new.replace(old, repl)
        if new != original:
            for run in para.runs:
                run.text = ""
            para.runs[0].text = new


def h1(doc, text):
    para = doc.add_paragraph(style="Heading 1")
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(14)


def h2(doc, text):
    para = doc.add_paragraph(style="Heading 2")
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(13)


def h3(doc, text):
    para = doc.add_paragraph(style="Heading 3")
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(12)


def p(doc, text, label=None):
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(4)
    para.paragraph_format.line_spacing = 1.08
    if label and text.startswith(label):
        run = para.add_run(label)
        run.bold = True
        para.add_run(text[len(label):])
    else:
        para.add_run(text)


def bullet(doc, text):
    para = doc.add_paragraph(style="List Bullet")
    para.paragraph_format.space_after = Pt(2)
    para.add_run(text)


def table(doc, headers, rows):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, header in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        cell.text = header
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for row in rows:
        cells = tbl.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
    doc.add_paragraph()


def already_added(doc: Document) -> bool:
    return any("PHẦN 8. CHUẨN HÓA CÁC Ý KHÓ VÀ CÔNG THỨC ĐINH" in p.text for p in doc.paragraphs)


def add_final_topics(doc: Document):
    doc.add_page_break()
    h1(doc, "PHẦN 8. CHUẨN HÓA CÁC Ý KHÓ VÀ CÔNG THỨC ĐINH")
    p(doc, "Phần này bổ sung các ý dễ bị hỏi theo kiểu khắt khe và trình bày bằng văn phong chuẩn để có thể đưa trực tiếp vào bài tự luận. Các mã tra cứu nhanh gồm [CH2-DOMINANCE], [CH3-OPERATIONALITY], [CH3-ANALYTICAL], [CH5-BV-DECOMP], [CH5-KNN-BV], [CH5-FORMULA].")

    h2(doc, "1. [CH2-DOMINANCE] Heuristic có thông tin nhiều hơn và hiệu quả của A*")
    p(doc, "Định nghĩa dominance: Một heuristic h2 được gọi là dominates h1 nếu với mọi nút n, h2(n) >= h1(n), đồng thời cả h1 và h2 đều admissible. Khi đó h2 gần chi phí thật hơn h1 nhưng vẫn không vượt quá chi phí thật.")
    p(doc, "Ý nghĩa đối với A*: A* dùng f(n)=g(n)+h(n). Nếu h lớn hơn nhưng vẫn admissible, f(n) của các nút không triển vọng sẽ tăng lên, làm A* ít phải mở rộng các nhánh kém. Vì vậy heuristic dominates thường giúp A* mở rộng ít nút hơn và chạy hiệu quả hơn.")
    p(doc, "Điều kiện cần lưu ý: Heuristic có thông tin nhiều hơn chỉ tốt khi vẫn giữ admissible hoặc consistent. Nếu h(n) vượt quá chi phí thật, A* có thể mất tính tối ưu.")
    p(doc, "Đoạn trả lời mẫu: Tính tối ưu của A* đến từ việc heuristic admissible làm f(n) trở thành cận dưới của chi phí lời giải đi qua n. Trong graph search, consistent quan trọng hơn vì nó đảm bảo f không giảm dọc đường đi, nên khi một nút đã được đóng thì không cần mở lại. Nếu một heuristic h2 dominates h1 và vẫn admissible, A* với h2 thường hiệu quả hơn vì loại bỏ được nhiều nút ít triển vọng hơn.")

    h2(doc, "2. [CH3-OPERATIONALITY] Operationality trong EBL")
    p(doc, "Khái niệm: Operationality trong EBL là yêu cầu rằng luật học được phải được biểu diễn bằng những điều kiện có thể kiểm tra hoặc thực thi trực tiếp bởi hệ thống. EBL không chỉ chứng minh một ví dụ đúng, mà còn biến tri thức trừu tượng thành quy tắc nhận diện nhanh các trường hợp tương tự.")
    p(doc, "Ví dụ: Một domain theory có thể nói rằng một vật có thể nhấc lên nếu lực cần thiết nhỏ hơn sức nâng của robot. Điều kiện này có thể quá trừu tượng nếu hệ thống không đo trực tiếp lực cần thiết. Một luật operational hơn có thể dùng các thuộc tính đo được như Light(x), HasHandle(x), Cup(x).")
    p(doc, "Liên hệ với Weakest Preimage: Weakest preimage giúp loại bỏ chi tiết thừa và giữ lại các điều kiện đủ để kết luận. Operationality bổ sung yêu cầu rằng các điều kiện được giữ lại phải dùng được trong nhận diện thực tế.")
    p(doc, "Ý kết luận: Một lời giải thích tốt trong EBL phải vừa đúng về mặt logic, vừa operational để cải thiện tốc độ và khả năng áp dụng của hệ thống trong các tình huống mới.")

    h2(doc, "3. [CH3-ANALYTICAL] Học phân tích, học quy nạp và khi nào EBL hơn Decision Tree")
    table(doc, ["Tiêu chí", "Analytical Learning/EBL", "Inductive Learning/Decision Tree"], [
        ("Nguồn tri thức", "Domain theory và một hoặc ít ví dụ", "Nhiều ví dụ huấn luyện"),
        ("Cơ chế", "Dùng suy diễn để giải thích ví dụ rồi tổng quát hóa proof", "Dùng dữ liệu để tìm thuộc tính chia tốt như entropy/gain"),
        ("Điểm mạnh", "Cần ít dữ liệu nếu tri thức nền đúng; luật có giải thích logic", "Ít phụ thuộc tri thức nền; học tốt khi có nhiều dữ liệu đại diện"),
        ("Điểm yếu", "Phụ thuộc mạnh vào domain theory; tri thức sai dẫn đến luật sai", "Dễ overfit nếu cây sâu; cần đủ dữ liệu"),
        ("Khi nên dùng", "Miền có luật nền rõ ràng, dữ liệu ít, cần giải thích", "Miền thiếu tri thức nền nhưng có dữ liệu nhãn đủ lớn"),
    ])
    p(doc, "Đoạn trả lời mẫu: EBL nên được dùng thay Decision Tree khi miền bài toán có domain theory chính xác và số lượng ví dụ huấn luyện ít. Trong trường hợp đó, EBL có thể dùng suy diễn để tạo luật tổng quát từ một ví dụ, giúp giảm sample complexity. Decision Tree phù hợp hơn khi tri thức nền yếu nhưng có nhiều dữ liệu đủ đại diện.")

    h2(doc, "4. [CH5-BV-DECOMP] Công thức phân rã Bias-Variance trong hồi quy")
    p(doc, "Công thức chuẩn: E[(y0 - f_hat(x0))^2] = Var(f_hat(x0)) + [Bias(f_hat(x0))]^2 + Var(epsilon).")
    p(doc, "Ý nghĩa từng thành phần: Var(f_hat(x0)) đo độ nhạy của mô hình với tập huấn luyện; Bias(f_hat(x0)) đo sai lệch hệ thống giữa dự đoán trung bình của mô hình và hàm thật; Var(epsilon) là nhiễu không thể giảm, còn gọi là irreducible error.")
    p(doc, "Giải thích quan trọng: Một mô hình quá đơn giản thường có bias cao và variance thấp. Một mô hình quá phức tạp thường có bias thấp nhưng variance cao. Mục tiêu của học máy là chọn độ phức tạp sao cho tổng ba thành phần, đặc biệt bias^2 và variance, được cân bằng.")
    p(doc, "Ý kết luận: Var(epsilon) không thể bị loại bỏ bằng cách đổi mô hình vì nó đến từ nhiễu nội tại của dữ liệu hoặc yếu tố không quan sát được. Mô hình chỉ có thể tác động đến bias và variance.")

    h2(doc, "5. [CH5-KNN-BV] Phân tích Bias-Variance qua KNN khi thay đổi K")
    p(doc, "K nhỏ: Khi K=1 hoặc rất nhỏ, KNN có ranh giới quyết định rất linh hoạt. Mô hình có bias thấp vì có thể khớp các mẫu huấn luyện phức tạp, nhưng variance cao vì chỉ cần dữ liệu thay đổi nhẹ thì láng giềng gần nhất và dự đoán có thể đổi.")
    p(doc, "K lớn: Khi K lớn, dự đoán được làm mượt bởi nhiều láng giềng. Variance giảm vì mô hình ít nhạy với từng điểm riêng lẻ, nhưng bias tăng vì ranh giới quyết định trở nên quá đơn giản và có thể bỏ qua cấu trúc cục bộ.")
    p(doc, "Kết luận chuẩn: Giảm bias thường làm tăng variance vì mô hình phải linh hoạt hơn để khớp dữ liệu. Với KNN, chọn K nhỏ làm mô hình linh hoạt và dễ overfit; chọn K lớn làm mô hình ổn định hơn nhưng có nguy cơ underfit. K nên được chọn bằng validation hoặc cross-validation.")

    h2(doc, "6. [CH5-FORMULA] Bảng công thức toán học nên có trong bài")
    table(doc, ["Chủ đề", "Công thức", "Ghi chú trình bày"], [
        ("Entropy", "H(S) = - sum_i p_i log2(p_i)", "Đo độ hỗn loạn của tập dữ liệu"),
        ("Information Gain", "Gain(S,A)=H(S)-sum_v |S_v|/|S| H(S_v)", "Mức giảm entropy khi chia theo thuộc tính A"),
        ("Gini", "Gini(S)=1-sum_i p_i^2", "Thường dùng trong CART"),
        ("Bias-Variance", "E[(y0-f_hat(x0))^2]=Var(f_hat)+Bias(f_hat)^2+Var(epsilon)", "Var(epsilon) là sai số không thể giảm"),
        ("Sample Error", "error_S(h)=số mẫu bị phân loại sai / |S|", "Đo trên tập mẫu hữu hạn"),
        ("True Error", "error_D(h)=Pr_{x~D}[h(x) != c(x)]", "Xác suất lỗi trên phân phối/quần thể dữ liệu"),
        ("Confidence Interval", "error_S ± z_N sqrt(error_S(1-error_S)/n)", "Ước lượng độ bất định của true error"),
        ("PAC hữu hạn", "m tăng theo (1/epsilon)(ln|H|+ln(1/delta))", "H lớn hoặc yêu cầu chính xác cao thì cần nhiều mẫu"),
    ])

    h2(doc, "7. Mối liên hệ liên chương ở mức nghiên cứu sâu")
    p(doc, "Chương 2 cung cấp cách nhìn học máy như tìm kiếm trong không gian giả thuyết. Chương 3 cung cấp domain theory và suy diễn logic. Chương 5 cung cấp phương pháp học từ dữ liệu và đánh giá tổng quát hóa. Khi kết hợp, tri thức nền ở Chương 3 có thể làm giảm số lượng dữ liệu cần thiết ở Chương 5.")
    p(doc, "Ví dụ: Trong EBL, nếu hệ đã có domain theory đúng, nó không cần hàng trăm ví dụ để học một quy tắc; một ví dụ được giải thích tốt có thể tạo ra luật tổng quát. Điều này thể hiện mối liên hệ giữa suy diễn logic và giảm sample complexity trong học máy.")
    p(doc, "Đoạn kết luận mẫu: Một hệ AI mạnh không chỉ tìm kiếm lời giải, không chỉ suy diễn từ luật, và không chỉ học thống kê từ dữ liệu. Sức mạnh đến từ việc kết hợp tìm kiếm, tri thức nền và học máy để vừa hiệu quả, vừa có khả năng giải thích, vừa tổng quát hóa tốt.")


def main():
    if not TARGET.exists():
        raise FileNotFoundError(TARGET)
    if not BACKUP.exists():
        copy2(TARGET, BACKUP)

    doc = Document(TARGET)
    normalize_font(doc)
    replace_text_preserve_runs(doc)
    if not already_added(doc):
        add_final_topics(doc)

    try:
        doc.save(TARGET)
        print(f"UPDATED={TARGET}")
    except PermissionError:
        doc.save(OUTPUT_IF_LOCKED)
        print(f"TARGET_LOCKED={TARGET}")
        print(f"OUTPUT={OUTPUT_IF_LOCKED}")
    print(f"BACKUP={BACKUP}")


if __name__ == "__main__":
    main()
