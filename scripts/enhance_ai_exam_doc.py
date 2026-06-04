from pathlib import Path
from shutil import copy2

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


TARGET = Path(r"C:\Users\dattr\Downloads\De_Cuong_AI_Tu_Luan_Tong_Hop_DA_TRA_LOI_CHI_TIET.docx")
BACKUP = TARGET.with_name(TARGET.stem + "_backup_truoc_khi_them_meo_thi.docx")


def normalize_font(doc: Document):
    for style_name in ["Normal", "Heading 1", "Heading 2", "Heading 3"]:
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        if style_name == "Normal":
            style.font.size = Pt(12)


def h1(doc, text):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(10)
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(14)


def h2(doc, text):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(8)
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(13)


def p(doc, text, bold_label=None):
    para = doc.add_paragraph()
    para.paragraph_format.first_line_indent = Inches(0.25)
    para.paragraph_format.space_after = Pt(4)
    para.paragraph_format.line_spacing = 1.08
    if bold_label and text.startswith(bold_label):
        run = para.add_run(bold_label)
        run.bold = True
        para.add_run(text[len(bold_label):])
    else:
        para.add_run(text)


def bullet(doc, text):
    para = doc.add_paragraph(style="List Bullet")
    para.paragraph_format.space_after = Pt(2)
    para.add_run(text)


def table(doc, headers, rows):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, header in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        cell.text = header
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for row in rows:
        cells = tbl.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    doc.add_paragraph()


def already_enhanced(doc: Document) -> bool:
    return any("PHẦN 4. PHẦN NÂNG ĐIỂM KHI ĐI THI" in para.text for para in doc.paragraphs)


def enhance():
    if not TARGET.exists():
        raise FileNotFoundError(TARGET)
    if not BACKUP.exists():
        copy2(TARGET, BACKUP)

    doc = Document(TARGET)
    normalize_font(doc)
    if already_enhanced(doc):
        print("ALREADY_ENHANCED=1")
        print(f"TARGET={TARGET}")
        print(f"BACKUP={BACKUP}")
        return

    doc.add_page_break()
    h1(doc, "PHẦN 4. PHẦN NÂNG ĐIỂM KHI ĐI THI")
    p(doc, "Nếu chỉ học định nghĩa thì có thể làm được câu cơ bản, nhưng để bài tự luận chắc điểm cao cần biết so sánh, chọn thuật toán theo tình huống, đưa ví dụ và chứng minh ngắn. Phần này bổ sung đúng các ý đó.")

    h2(doc, "1. Khi đề yêu cầu so sánh, nên viết theo khung nào?")
    p(doc, "Khung trả lời: nêu điểm giống nhau, nêu tiêu chí so sánh, lập bảng hoặc đoạn so sánh, đưa ví dụ tình huống, cuối cùng kết luận nên chọn phương pháp nào trong điều kiện nào.", "Khung trả lời:")
    bullet(doc, "Điểm giống nhau: hai phương pháp cùng giải quyết bài toán gì.")
    bullet(doc, "Điểm khác nhau: cấu trúc dữ liệu, tiêu chí chọn bước tiếp theo, tính hoàn chỉnh, tính tối ưu, bộ nhớ, ứng dụng.")
    bullet(doc, "Ví dụ: dùng một bài toán tìm đường, phân loại email hoặc chẩn đoán bệnh để minh họa.")
    bullet(doc, "Kết luận: không nói phương pháp nào luôn tốt hơn, mà nói tốt hơn trong điều kiện cụ thể nào.")

    h2(doc, "2. Các cặp so sánh rất dễ ra thi")
    table(doc, ["Cặp so sánh", "Giống nhau", "Khác nhau cốt lõi", "Kết luận nên viết"], [
        ("BFS và DFS", "Đều là tìm kiếm không có thông tin", "BFS duyệt theo mức bằng queue; DFS đi sâu bằng stack/đệ quy", "BFS phù hợp khi cần lời giải nông nhất; DFS phù hợp khi bộ nhớ hạn chế"),
        ("UCS và BFS", "Đều có thể tối ưu trong điều kiện phù hợp", "BFS tối ưu theo số bước khi chi phí bằng nhau; UCS tối ưu theo tổng chi phí g(n)", "Nếu cạnh có trọng số khác nhau thì chọn UCS"),
        ("Greedy và A*", "Đều dùng heuristic", "Greedy chỉ dùng h(n); A* dùng g(n)+h(n)", "Greedy nhanh nhưng dễ sai; A* đáng tin hơn nếu heuristic admissible/consistent"),
        ("Forward và Backward Chaining", "Đều suy diễn từ fact và rule", "Forward đi từ dữ kiện; Backward đi từ mục tiêu", "Forward hợp khi muốn suy ra nhiều kết luận; Backward hợp khi có truy vấn cụ thể"),
        ("Logic mệnh đề và logic vị từ", "Đều biểu diễn tri thức bằng logic", "Logic vị từ có đối tượng, quan hệ, biến và lượng từ", "Logic vị từ mạnh hơn nhưng suy diễn phức tạp hơn"),
        ("Overfitting và Underfitting", "Đều làm mô hình tổng quát hóa kém", "Overfit: train tốt test kém; underfit: train và test đều kém", "Cần cân bằng bias-variance, dùng validation để chọn mô hình"),
        ("KNN và Logistic Regression", "Đều dùng cho phân loại", "KNN dựa vào láng giềng; Logistic học tham số tuyến tính xác suất", "KNN linh hoạt nhưng dự đoán chậm; Logistic nhanh và dễ giải thích"),
        ("Decision Tree và Neural Network", "Đều học quan hệ từ dữ liệu", "Tree dễ giải thích; NN học phi tuyến phức tạp nhưng khó giải thích", "Tree hợp bài toán cần giải thích; NN hợp dữ liệu lớn/phức tạp"),
    ])

    h2(doc, "3. Tình huống chọn thuật toán trong bài thi")
    table(doc, ["Tình huống đề cho", "Nên chọn/nhắc tới", "Lý do"], [
        ("Tìm đường không trọng số, cần ít bước nhất", "BFS", "BFS tìm lời giải nông nhất và tối ưu khi chi phí bước bằng nhau"),
        ("Tìm đường có trọng số km/thời gian/chi phí", "UCS hoặc A*", "UCS tối ưu theo g(n); A* nhanh hơn nếu có heuristic tốt"),
        ("Có khoảng cách ước lượng đến đích", "Greedy hoặc A*", "Greedy dùng h(n), A* kết hợp g(n) và h(n) nên thường an toàn hơn"),
        ("Bộ nhớ rất hạn chế", "DFS hoặc IDDFS", "DFS/IDDFS dùng ít bộ nhớ hơn BFS/A*"),
        ("Cần chứng minh một truy vấn cụ thể trong KB", "Backward chaining", "Đi từ mục tiêu nên tránh suy diễn lan man"),
        ("Có nhiều dữ kiện cảm biến liên tục đổ vào", "Forward chaining", "Đi từ dữ kiện để tự động sinh cảnh báo/kết luận"),
        ("Dữ liệu có nhãn, cần dự đoán y", "Supervised learning", "Học ánh xạ từ x sang y"),
        ("Dữ liệu không nhãn, cần phân nhóm", "Unsupervised learning/K-means", "Tìm cấu trúc ẩn trong dữ liệu"),
        ("Agent học qua thưởng phạt", "Reinforcement learning/Q-learning", "Tối đa hóa tổng reward dài hạn"),
    ])

    h2(doc, "4. Ví dụ chứng minh/ngắn gọn nên đưa vào bài")
    p(doc, "Ví dụ 1 - Vì sao BFS tối ưu khi chi phí bằng nhau: BFS mở rộng toàn bộ node ở độ sâu k trước khi sang độ sâu k+1. Nếu mỗi bước có cùng chi phí, lời giải có độ sâu nhỏ hơn luôn có chi phí nhỏ hơn hoặc bằng lời giải sâu hơn. Vì vậy lời giải đầu tiên BFS tìm thấy là lời giải ít bước nhất.", "Ví dụ 1")
    p(doc, "Ví dụ 2 - Vì sao Greedy có thể không tối ưu: giả sử từ A đi đến B có h(B)=1 nhưng chi phí A-B rất lớn, còn đi đến C có h(C)=5 nhưng sau đó đến đích rất rẻ. Greedy chọn B vì h nhỏ hơn, nhưng tổng đường đi qua B có thể đắt hơn đường qua C. Lý do là Greedy bỏ qua g(n).", "Ví dụ 2")
    p(doc, "Ví dụ 3 - Vì sao A* cần admissible heuristic: nếu h(n) đánh giá quá cao chi phí thật, A* có thể bỏ qua một nhánh thực sự tối ưu vì tưởng nhánh đó đắt. Khi h(n) không vượt quá chi phí thật, A* không loại nhầm đường tối ưu.", "Ví dụ 3")
    p(doc, "Ví dụ 4 - Entropy và Information Gain: nếu một tập có 10 mẫu gồm 5 dương và 5 âm thì entropy cao vì dữ liệu lẫn nhiều. Nếu chia theo thuộc tính A tạo hai nhóm mỗi nhóm gần như cùng lớp, entropy sau chia giảm mạnh, nên information gain của A lớn.", "Ví dụ 4")
    p(doc, "Ví dụ 5 - Bayes trong Naive Bayes: khi phân loại email, ta so sánh P(Spam|words) và P(NotSpam|words). Naive Bayes giả định các từ độc lập có điều kiện theo lớp để tính nhanh tích các xác suất P(word|class).", "Ví dụ 5")
    p(doc, "Ví dụ 6 - Bias và variance: dùng đường thẳng cho dữ liệu dạng cong là bias cao vì mô hình quá đơn giản. Dùng đa thức bậc rất cao đi qua gần mọi điểm train là variance cao vì mô hình học cả nhiễu.", "Ví dụ 6")

    h2(doc, "5. Các câu hỏi biến dạng và cách bẻ về bài đã học")
    table(doc, ["Đề hỏi kiểu lạ", "Cách nhận diện", "Bẻ về nội dung cần viết"], [
        ("Thuật toán nào phù hợp hơn trong bài toán tìm đường?", "Có/không có trọng số, có/không có heuristic", "So sánh BFS, UCS, Greedy, A* theo cost và heuristic"),
        ("Vì sao mô hình dự đoán tốt trên train nhưng kém trên test?", "Train tốt test kém", "Overfitting, variance cao, regularization, cross-validation"),
        ("Vì sao mô hình quá đơn giản dự đoán kém?", "Train và test đều kém", "Underfitting, bias cao, tăng độ phức tạp/thêm đặc trưng"),
        ("Khi nào dùng logic vị từ thay logic mệnh đề?", "Có đối tượng, quan hệ, lượng từ", "Nêu hạn chế logic mệnh đề và ưu điểm FOL"),
        ("Hệ chuyên gia suy ra kết luận như thế nào?", "Có fact và rule", "KB, inference, forward/backward chaining"),
        ("Học máy khác lập trình truyền thống thế nào?", "Có dữ liệu/kinh nghiệm", "Định nghĩa T, P, E của Mitchell"),
    ])

    h2(doc, "6. Mẫu mở bài và kết luận dùng được cho nhiều câu")
    p(doc, "Mẫu mở bài tìm kiếm: Trong AI, tìm kiếm là cách mô hình hóa bài toán thành không gian trạng thái, trong đó tác tử cần tìm chuỗi hành động từ trạng thái ban đầu đến trạng thái đích. Sự khác nhau giữa các thuật toán nằm ở cách chọn nút tiếp theo để mở rộng.", "Mẫu mở bài tìm kiếm:")
    p(doc, "Mẫu kết luận tìm kiếm: Không có thuật toán tìm kiếm tốt nhất cho mọi trường hợp. Nếu chi phí bằng nhau có thể dùng BFS; nếu chi phí khác nhau nên dùng UCS; nếu có heuristic tốt thì A* thường hiệu quả và vẫn đảm bảo tối ưu khi heuristic thỏa điều kiện.", "Mẫu kết luận tìm kiếm:")
    p(doc, "Mẫu mở bài logic: Biểu diễn tri thức và suy diễn giúp hệ AI không chỉ lưu dữ kiện mà còn rút ra kết luận mới từ các luật đã biết. Đây là nền tảng của hệ chuyên gia và các hệ thống AI dựa trên tri thức.", "Mẫu mở bài logic:")
    p(doc, "Mẫu kết luận logic: Logic mệnh đề đơn giản và dễ xử lý, còn logic vị từ biểu diễn mạnh hơn nhờ đối tượng, quan hệ và lượng từ. Việc chọn phương pháp suy diễn phụ thuộc vào việc ta muốn suy ra tất cả kết luận hay chỉ chứng minh một mục tiêu cụ thể.", "Mẫu kết luận logic:")
    p(doc, "Mẫu mở bài học máy: Học máy là cách làm cho chương trình cải thiện hiệu năng nhờ kinh nghiệm thay vì chỉ dựa vào quy tắc lập trình cứng. Theo Mitchell, cần xác định rõ nhiệm vụ T, thước đo P và kinh nghiệm E.", "Mẫu mở bài học máy:")
    p(doc, "Mẫu kết luận học máy: Mục tiêu quan trọng nhất của học máy không phải là học thuộc training set mà là tổng quát hóa tốt trên dữ liệu mới. Vì vậy cần quan tâm đến train/validation/test, overfitting, regularization và cross-validation.", "Mẫu kết luận học máy:")

    h2(doc, "7. Những lỗi dễ mất điểm")
    bullet(doc, "Chỉ nêu định nghĩa mà không có ví dụ. Bài tự luận nên luôn có ít nhất một ví dụ cụ thể.")
    bullet(doc, "Nhầm Greedy với A*: Greedy dùng h(n), còn A* dùng g(n)+h(n).")
    bullet(doc, "Nói DFS tối ưu. DFS thường không tối ưu và có thể không hoàn chỉnh nếu không gian vô hạn.")
    bullet(doc, "Nói BFS luôn tối ưu. BFS chỉ tối ưu khi chi phí mỗi bước bằng nhau.")
    bullet(doc, "Nhầm overfitting và underfitting: overfit là train tốt test kém; underfit là cả train và test đều kém.")
    bullet(doc, "Chỉ nói accuracy khi dữ liệu mất cân bằng. Nên nhắc thêm precision, recall và F1.")
    bullet(doc, "Viết công thức nhưng không giải thích ký hiệu. Khi ghi f(n)=g(n)+h(n), phải nói rõ g và h là gì.")
    bullet(doc, "So sánh nhưng không kết luận nên dùng cái nào trong tình huống nào.")

    h2(doc, "8. Nếu thầy ra câu tổng hợp lớn, nên triển khai như sau")
    p(doc, "Câu tổng hợp Chương 2: mở bài về state space, nêu các thành phần bài toán tìm kiếm, phân loại uninformed/informed search, so sánh BFS-DFS-UCS-Greedy-A*, đưa ví dụ tìm đường, kết luận chọn thuật toán theo chi phí và heuristic.", "Câu tổng hợp Chương 2:")
    p(doc, "Câu tổng hợp Chương 3: mở bài về vai trò tri thức, nêu KB gồm fact/rule, so sánh logic mệnh đề và logic vị từ, trình bày forward/backward chaining, resolution/unification, kết luận ứng dụng trong hệ chuyên gia.", "Câu tổng hợp Chương 3:")
    p(doc, "Câu tổng hợp Chương 5: mở bài bằng định nghĩa Mitchell, phân loại supervised/unsupervised/reinforcement, trình bày vài thuật toán tiêu biểu, nói về đánh giá mô hình, overfitting-underfitting, bias-variance, regularization và cross-validation.", "Câu tổng hợp Chương 5:")

    doc.save(TARGET)
    print(f"UPDATED={TARGET}")
    print(f"BACKUP={BACKUP}")


if __name__ == "__main__":
    enhance()
