from pathlib import Path
from shutil import copy2

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Pt


TARGET = Path(r"D:\GK_AI\CuoiKiAI.docx")
BACKUP = TARGET.with_name("CuoiKiAI_backup_truoc_khi_bo_sung_khac_khe.docx")
OUTPUT_IF_LOCKED = TARGET.with_name("CuoiKiAI_DA_BO_SUNG_KHAC_KHE.docx")


def normalize_font(doc: Document):
    for style_name in ["Normal", "Heading 1", "Heading 2", "Heading 3"]:
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        if style_name == "Normal":
            style.font.size = Pt(12)


def h1(doc, text):
    para = doc.add_paragraph(style="Heading 1")
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(14)


def h2(doc, text):
    para = doc.add_paragraph(style="Heading 2")
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(13)


def h3(doc, text):
    para = doc.add_paragraph(style="Heading 3")
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(12)


def p(doc, text, label=None):
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(4)
    para.paragraph_format.line_spacing = 1.08
    if label and text.startswith(label):
        run = para.add_run(label)
        run.bold = True
        para.add_run(text[len(label):])
    else:
        para.add_run(text)


def bullet(doc, text):
    para = doc.add_paragraph(style="List Bullet")
    para.paragraph_format.space_after = Pt(2)
    para.add_run(text)


def table(doc, headers, rows):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, header in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        cell.text = header
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for row in rows:
        cells = tbl.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    doc.add_paragraph()


def already_added(doc):
    return any("PHẦN 7. BỔ SUNG KHẮC KHE ĐỂ LẤY ĐIỂM TUYỆT ĐỐI" in x.text for x in doc.paragraphs)


def add_section(doc):
    doc.add_page_break()
    h1(doc, "PHẦN 7. BỔ SUNG KHẮC KHE ĐỂ LẤY ĐIỂM TUYỆT ĐỐI")
    p(doc, "Phần này bổ sung các câu hỏi khó, công thức và ví dụ phản biện để xử lý đề mở ở mức nâng cao. Dùng Ctrl+F các mã: [CH2-ML-AS-SEARCH], [CH2-ASTAR-PROOF], [CH2-GENERALITY], [CH3-EBL-VS-INDUCTIVE], [CH3-HORN-MGU], [CH5-GINI], [CH5-KERNEL], [CH5-PAC], [CH5-LAZY-EAGER].")

    h2(doc, "1. [CH2-ML-AS-SEARCH] Tại sao học máy được coi là bài toán tìm kiếm?")
    p(doc, "Ý chính: Trong nhiều mô hình học máy, học có thể xem là quá trình tìm kiếm một giả thuyết h tốt nhất trong không gian giả thuyết H. Dữ liệu huấn luyện đóng vai trò ràng buộc hoặc tín hiệu đánh giá, còn thuật toán học là chiến lược tìm kiếm trong H.")
    p(doc, "Liên hệ với concept learning: Với học khái niệm, H là tập mọi mô tả khái niệm có thể. Thuật toán như Find-S hoặc Candidate-Elimination tìm các giả thuyết nhất quán với ví dụ dương/âm. Version Space chính là vùng chưa bị loại trong không gian tìm kiếm.")
    p(doc, "Liên hệ với học máy hiện đại: Linear regression tìm vector theta làm loss nhỏ nhất; decision tree tìm cấu trúc cây tốt; neural network tìm bộ trọng số tối ưu. Dù cách tìm khác nhau, bản chất vẫn là tìm trong một không gian mô hình/tham số rất lớn.")
    p(doc, "Câu kết luận nên chép: Học máy là tìm kiếm có định hướng bởi dữ liệu, trong đó mục tiêu không chỉ là khớp training set mà là tìm giả thuyết tổng quát hóa tốt trên dữ liệu chưa thấy.")

    h2(doc, "2. [CH2-ASTAR-PROOF] Chứng minh ngắn tính tối ưu của A* và vai trò Consistent")
    p(doc, "A* dùng f(n)=g(n)+h(n). Nếu h admissible, tức h(n) không vượt quá chi phí thật từ n đến goal, thì f(n) là cận dưới của chi phí lời giải đi qua n. Vì vậy A* không đánh giá một đường tối ưu là đắt hơn thực tế.")
    p(doc, "Sketch chứng minh: Gọi C* là chi phí lời giải tối ưu. Trên đường tối ưu luôn tồn tại một nút n chưa mở rộng có f(n)=g(n)+h(n) <= C* do h admissible. Nếu A* chọn một goal G không tối ưu với cost g(G)>C*, thì vẫn còn nút n trên đường tối ưu có f(n)<=C*<g(G). A* lẽ ra phải chọn n trước G, mâu thuẫn. Vậy goal đầu tiên A* chọn là tối ưu.")
    p(doc, "Vì sao consistent quan trọng hơn trong graph search: Admissible chỉ đảm bảo không nói quá so với goal. Consistent còn đảm bảo h(n) <= c(n,a,n') + h(n'), làm f-value không giảm dọc đường đi. Khi heuristic consistent, lần đầu một node được lấy ra khỏi priority queue thì đường tới node đó đã tối ưu, nên graph search không cần mở lại node.")
    p(doc, "Câu chốt: Admissible đủ quan trọng cho optimality trong tree search; consistent mạnh hơn và thực tế hơn cho graph search vì xử lý trạng thái lặp an toàn hơn.")

    h2(doc, "3. [CH2-GENERALITY] Quan hệ tổng quát hóa >=g trong Concept Learning")
    p(doc, "Định nghĩa toán học: Giả thuyết h_i được gọi là tổng quát hơn hoặc bằng h_j, ký hiệu h_i >=g h_j, nếu mọi ví dụ được h_j phân loại là dương thì cũng được h_i phân loại là dương. Nói cách khác, tập ví dụ dương mà h_j bao phủ là tập con của tập ví dụ dương mà h_i bao phủ.")
    p(doc, "Ký hiệu theo tập: h_i >=g h_j khi PositiveSet(h_j) subset PositiveSet(h_i). Nếu h_i bao phủ nhiều trường hợp hơn, h_i tổng quát hơn; nếu bao phủ ít hơn, h_i cụ thể hơn.")
    p(doc, "Ví dụ EnjoySport: Giả thuyết <?, Warm, ?, ?, ?, ?> tổng quát hơn <Sunny, Warm, ?, ?, ?, ?> vì mọi ví dụ có Sky=Sunny và AirTemp=Warm chắc chắn cũng thỏa AirTemp=Warm, nhưng chiều ngược lại không đúng.")
    p(doc, "Vai trò trong Candidate-Elimination: S là biên cụ thể nhất, G là biên tổng quát nhất. Version Space gồm các giả thuyết nằm giữa hai biên theo quan hệ >=g.")

    h2(doc, "4. [CH3-EBL-VS-INDUCTIVE] So sánh EBL và Inductive Learning")
    table(doc, ["Tiêu chí", "Explanation-Based Learning", "Inductive Learning"], [
        ("Nguồn học", "Tri thức nền/domain theory + một hoặc ít ví dụ", "Nhiều ví dụ huấn luyện"),
        ("Cơ chế", "Giải thích vì sao ví dụ đúng rồi tổng quát hóa proof", "Tìm quy luật thống kê/giả thuyết khớp dữ liệu"),
        ("Kết luận", "Thường mang tính suy diễn nếu domain theory đúng", "Mang tính quy nạp, có thể sai trên dữ liệu mới"),
        ("Khi vượt trội", "Khi domain theory mạnh, đúng và dữ liệu ít", "Khi domain theory yếu nhưng có nhiều dữ liệu"),
        ("Rủi ro", "Phụ thuộc vào tri thức nền", "Cần dữ liệu đủ đại diện, dễ overfit/underfit"),
    ])
    p(doc, "Câu trả lời mẫu: EBL vượt trội khi ta có domain theory chính xác vì nó học được quy tắc tổng quát từ rất ít ví dụ, làm giảm sample complexity. Inductive learning phù hợp hơn khi thiếu tri thức nền nhưng có nhiều dữ liệu quan sát.")

    h2(doc, "5. [CH3-HORN-MGU] Horn Clause, Unification và MGU")
    p(doc, "Horn clause: Là mệnh đề logic có tối đa một literal dương. Dạng thường gặp trong luật là A1 and A2 and ... and Ak -> B, tương đương not A1 or not A2 or ... or B. Horn clauses là nền tảng của Prolog, forward/backward chaining và nhiều hệ ILP.")
    p(doc, "Vì sao Horn clause quan trọng: Nó cân bằng giữa khả năng biểu diễn và hiệu quả suy diễn. Với Horn clauses, forward chaining và backward chaining có thể suy luận hiệu quả hơn so với logic vị từ tổng quát.")
    p(doc, "Unification: Là quá trình tìm phép thế biến để hai biểu thức trở nên giống nhau. Ví dụ Parent(x,y) và Parent(John,Mary) hợp nhất bằng {x/John, y/Mary}.")
    p(doc, "MGU - Most General Unifier: Là phép hợp nhất tổng quát nhất, không ràng buộc biến nhiều hơn cần thiết. MGU giúp luật tổng quát áp dụng cho dữ kiện cụ thể mà vẫn giữ được tính tổng quát cho suy luận tiếp theo.")
    p(doc, "Ví dụ MGU: Loves(x,y) và Loves(An,z) có MGU {x/An, y/z}. Nếu thay luôn y/Binh thì vẫn hợp nhất được với một trường hợp cụ thể, nhưng kém tổng quát hơn MGU.")

    h2(doc, "6. [CH5-GINI] Entropy vs Gini trong Decision Tree")
    p(doc, "Entropy thường gắn với ID3/C4.5, đo độ hỗn loạn: H(S)=-sum p_i log2(p_i). Information Gain chọn thuộc tính làm giảm entropy nhiều nhất.")
    p(doc, "Gini impurity thường dùng trong CART: Gini(S)=1-sum p_i^2. Với hai lớp, Gini=1-p^2-(1-p)^2. Gini bằng 0 khi nút thuần nhất và cao khi các lớp trộn lẫn.")
    p(doc, "Ví dụ tính nhanh: Nếu một nút có 6 mẫu dương và 4 mẫu âm, p+=0.6, p-=0.4. Gini=1-0.6^2-0.4^2=1-0.36-0.16=0.48. Entropy=-0.6log2(0.6)-0.4log2(0.4)≈0.971.")
    p(doc, "So sánh: Entropy dùng log nên tính nặng hơn một chút và nhạy theo thông tin; Gini đơn giản hơn, thường cho kết quả tách khá giống entropy. Trong bài thi, nhắc ID3 dùng entropy/gain, CART hay dùng Gini.")

    h2(doc, "7. [CH5-KERNEL] Vì sao Kernel giúp SVM xử lý phi tuyến?")
    p(doc, "Vấn đề: Dữ liệu có thể không phân tách tuyến tính trong không gian gốc. Ý tưởng là ánh xạ dữ liệu sang không gian đặc trưng cao chiều hơn, nơi có thể tồn tại siêu phẳng tuyến tính phân tách dữ liệu.")
    p(doc, "Kernel trick: SVM chỉ cần tích vô hướng giữa các điểm trong không gian đặc trưng, không nhất thiết phải tính trực tiếp tọa độ sau ánh xạ. Kernel K(x,z)=phi(x).phi(z) cho phép tính tích vô hướng đó ngay trong không gian gốc.")
    p(doc, "Ví dụ trực giác: Các điểm nằm vòng trong/vòng ngoài không tách được bằng đường thẳng trong 2D. Nếu thêm đặc trưng r^2=x1^2+x2^2, dữ liệu có thể tách bằng một ngưỡng theo bán kính. Kernel cho phép làm điều tương tự ở không gian cao chiều mà không phải tự tạo toàn bộ đặc trưng.")
    p(doc, "Câu chốt: Kernel không làm SVM 'ma thuật' hơn, mà giúp SVM học ranh giới phi tuyến trong không gian gốc thông qua siêu phẳng tuyến tính ở không gian đặc trưng cao chiều.")

    h2(doc, "8. [CH5-CV-BIAS-VARIANCE] Bias-Variance khi thay đổi k trong k-fold CV")
    p(doc, "K nhỏ, ví dụ 5-fold: Mỗi lần train dùng khoảng 80% dữ liệu. Ước lượng test error có thể bias cao hơn một chút vì mô hình được train trên ít dữ liệu hơn toàn bộ tập, nhưng variance thường thấp hơn vì các fold validation lớn hơn và ít tương quan hơn.")
    p(doc, "K lớn, đặc biệt LOOCV với k=n: Mỗi lần train dùng n-1 mẫu nên bias thấp, vì gần giống train trên toàn bộ dữ liệu. Tuy nhiên variance có thể cao hơn vì các tập train gần như giống nhau và mỗi validation set chỉ có một điểm, làm ước lượng nhạy với từng quan sát.")
    p(doc, "Câu kết luận: Trong thực tế 5-fold hoặc 10-fold thường là lựa chọn cân bằng giữa chi phí tính toán, bias và variance. LOOCV hữu ích khi dữ liệu rất ít nhưng có thể tốn kém.")

    h2(doc, "9. [CH5-IG-EXAMPLE] Ví dụ tính Information Gain kiểu PlayTennis")
    p(doc, "Giả sử tập S có 14 mẫu: 9 Yes và 5 No. Entropy(S)=-(9/14)log2(9/14)-(5/14)log2(5/14)≈0.940.")
    p(doc, "Xét thuộc tính Outlook có 3 giá trị: Sunny gồm 5 mẫu (2 Yes, 3 No), Overcast gồm 4 mẫu (4 Yes, 0 No), Rain gồm 5 mẫu (3 Yes, 2 No).")
    p(doc, "Entropy(Sunny)=-(2/5)log2(2/5)-(3/5)log2(3/5)≈0.971. Entropy(Overcast)=0 vì toàn Yes. Entropy(Rain)=-(3/5)log2(3/5)-(2/5)log2(2/5)≈0.971.")
    p(doc, "Entropy sau khi chia theo Outlook = (5/14)*0.971 + (4/14)*0 + (5/14)*0.971 ≈ 0.694. Information Gain = 0.940 - 0.694 = 0.246.")
    p(doc, "Cách trình bày: Viết công thức tổng quát trước, thay số sau, rồi kết luận thuộc tính có gain lớn hơn sẽ được ưu tiên chọn tại node.")

    h2(doc, "10. [CH5-PAC] PAC Learning - Bao nhiêu dữ liệu là đủ?")
    p(doc, "PAC là Probably Approximately Correct. Ý tưởng: thuật toán học được xem là tốt nếu với xác suất cao, giả thuyết học được có sai số nhỏ. 'Probably' liên quan đến độ tin cậy 1-delta; 'Approximately Correct' liên quan đến sai số tối đa epsilon.")
    p(doc, "Diễn giải: Ta muốn tìm h sao cho error(h) <= epsilon với xác suất ít nhất 1-delta. Epsilon càng nhỏ thì yêu cầu chính xác càng cao; delta càng nhỏ thì yêu cầu tin cậy càng cao; cả hai làm số mẫu cần thiết tăng.")
    p(doc, "Công thức trực giác cho H hữu hạn: Số mẫu thường tăng theo (1/epsilon) * (ln|H| + ln(1/delta)). Không cần thuộc máy móc nếu đề không yêu cầu, nhưng nên hiểu: không gian giả thuyết càng lớn, muốn chính xác/tin cậy hơn thì cần nhiều dữ liệu hơn.")
    p(doc, "Liên hệ inductive bias: Nếu H nhỏ hơn do bias mạnh hơn, sample complexity có thể giảm. Nhưng H quá nhỏ có thể không chứa target concept, gây bias sai.")

    h2(doc, "11. [CH5-LAZY-EAGER] Lazy Learning vs Eager Learning")
    table(doc, ["Tiêu chí", "Lazy Learning", "Eager Learning"], [
        ("Ý tưởng", "Trì hoãn tổng quát hóa đến lúc có truy vấn", "Xây mô hình tổng quát ngay trong lúc huấn luyện"),
        ("Ví dụ", "KNN, case-based reasoning", "Decision Tree, Logistic Regression, Neural Network"),
        ("Huấn luyện", "Nhanh, chủ yếu lưu dữ liệu", "Chậm hơn vì phải học mô hình/tham số"),
        ("Dự đoán", "Chậm hơn vì phải so sánh với dữ liệu train", "Nhanh hơn vì đã có mô hình"),
        ("Ưu điểm", "Linh hoạt, dễ cập nhật dữ liệu mới", "Nén tri thức, dự đoán nhanh, dễ triển khai"),
        ("Nhược điểm", "Tốn bộ nhớ, nhạy với khoảng cách/thang đo", "Có thể mất thông tin nếu mô hình/bias không phù hợp"),
    ])
    p(doc, "Câu chốt: KNN là lazy vì gần như không học mô hình trước; Decision Tree/ANN là eager vì học cấu trúc hoặc trọng số trước khi dự đoán.")

    h2(doc, "12. [INTER-CHAPTER] Mối liên hệ liên chương: Chương 3 hỗ trợ Chương 5 thế nào?")
    p(doc, "Chương 3 cung cấp tri thức nền và cơ chế suy diễn; Chương 5 cung cấp khả năng học từ dữ liệu. Khi kết hợp, domain theory có thể giảm sample complexity vì mô hình không phải học mọi thứ từ đầu.")
    p(doc, "Ví dụ EBL: Nếu đã có luật vật nhẹ và có tay cầm thì dễ nhấc, hệ thống chỉ cần một ví dụ cái cốc nhẹ để tổng quát thành quy tắc Liftable(x), thay vì cần rất nhiều ví dụ thống kê.")
    p(doc, "Ví dụ ILP/FOIL: Dữ liệu quan hệ như Parent, Ancestor, Male, Female khó biểu diễn bằng bảng thuộc tính phẳng. Logic giúp biểu diễn quan hệ, còn học máy giúp học luật mới từ ví dụ.")
    p(doc, "Câu kết luận: Tri thức logic giúp học máy học nhanh hơn, giải thích tốt hơn và cần ít dữ liệu hơn; học máy giúp hệ logic thích nghi khi tri thức thủ công chưa đầy đủ.")

    h2(doc, "13. Bảng công thức nâng cao cần chép nhanh")
    table(doc, ["Mục", "Công thức/ý chính", "Dùng khi"], [
        ("Generality", "h_i >=g h_j nếu PositiveSet(h_j) subset PositiveSet(h_i)", "Candidate-Elimination"),
        ("A*", "f(n)=g(n)+h(n)", "Tìm kiếm heuristic"),
        ("Consistent", "h(n) <= c(n,a,n') + h(n')", "A* graph search"),
        ("Gini", "1 - sum p_i^2", "CART decision tree"),
        ("Entropy", "-sum p_i log2 p_i", "ID3/information gain"),
        ("FOIL-Gain", "t[log2(p1/(p1+n1))-log2(p0/(p0+n0))]", "Chọn literal trong FOIL"),
        ("CI Error", "error_S ± z sqrt(error_S(1-error_S)/n)", "Khoảng tin cậy sai số"),
        ("PAC trực giác", "m tăng theo (1/epsilon)(ln|H|+ln(1/delta))", "Bao nhiêu dữ liệu là đủ"),
    ])

    h2(doc, "14. Mẫu câu trả lời khó đã ghép sẵn")
    p(doc, "Mẫu 1 - Học máy là tìm kiếm: Học máy có thể xem là tìm kiếm trong không gian giả thuyết H. Mỗi giả thuyết là một mô hình có thể giải thích dữ liệu. Thuật toán học dùng dữ liệu để loại bỏ hoặc giảm điểm các giả thuyết kém, từ đó tìm h tối ưu theo tiêu chí như consistency, accuracy hoặc loss. Concept learning thể hiện rõ điều này qua version space, còn regression/neural network thể hiện qua tìm kiếm trong không gian tham số.")
    p(doc, "Mẫu 2 - EBL vs quy nạp: EBL học bằng cách giải thích một ví dụ dựa trên domain theory rồi tổng quát hóa proof, nên rất hiệu quả khi tri thức nền đúng và dữ liệu ít. Inductive learning học quy luật từ nhiều ví dụ nên phù hợp khi thiếu tri thức nền nhưng có dữ liệu đủ lớn. EBL có thể giảm sample complexity nhưng phụ thuộc mạnh vào chất lượng domain theory.")
    p(doc, "Mẫu 3 - SVM kernel: Kernel giúp SVM xử lý phi tuyến bằng cách ngầm ánh xạ dữ liệu sang không gian đặc trưng cao chiều, nơi có thể phân tách tuyến tính. Vì SVM chỉ cần tích vô hướng giữa các điểm, kernel K(x,z) cho phép tính tích vô hướng trong không gian mới mà không cần tính phi(x) trực tiếp. Đây là kernel trick.")


def main():
    if not TARGET.exists():
        raise FileNotFoundError(TARGET)
    if not BACKUP.exists():
        copy2(TARGET, BACKUP)
    doc = Document(TARGET)
    normalize_font(doc)
    if not already_added(doc):
        add_section(doc)
    try:
        doc.save(TARGET)
        print(f"UPDATED={TARGET}")
    except PermissionError:
        doc.save(OUTPUT_IF_LOCKED)
        print(f"TARGET_LOCKED={TARGET}")
        print(f"OUTPUT={OUTPUT_IF_LOCKED}")
    print(f"BACKUP={BACKUP}")


if __name__ == "__main__":
    main()
