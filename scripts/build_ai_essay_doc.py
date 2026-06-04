from pathlib import Path
from shutil import copy2

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


DOWNLOADS = Path(r"C:\Users\dattr\Downloads")
TARGET = DOWNLOADS / "De_Cuong_AI_Tu_Luan_Tong_Hop.docx"
OUTPUT = DOWNLOADS / "De_Cuong_AI_Tu_Luan_Tong_Hop_DA_TRA_LOI_CHI_TIET.docx"
BACKUP = DOWNLOADS / "De_Cuong_AI_Tu_Luan_Tong_Hop_backup_truoc_khi_bo_sung.docx"


def setup_doc() -> Document:
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.75)
    sec.bottom_margin = Inches(0.75)
    sec.left_margin = Inches(0.85)
    sec.right_margin = Inches(0.85)

    for style_name in ["Normal", "Title", "Heading 1", "Heading 2", "Heading 3"]:
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        if style_name == "Normal":
            style.font.size = Pt(12)
    return doc


def p(doc, text="", first_line=True, bold_label=None):
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(4)
    para.paragraph_format.line_spacing = 1.08
    if first_line:
        para.paragraph_format.first_line_indent = Inches(0.25)
    if bold_label and text.startswith(bold_label):
        run = para.add_run(bold_label)
        run.bold = True
        para.add_run(text[len(bold_label):])
    else:
        para.add_run(text)
    return para


def title(doc, text):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(16)


def h1(doc, text):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(8)
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(14)


def h2(doc, text):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(6)
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(13)


def bullet(doc, text):
    para = doc.add_paragraph(style="List Bullet")
    para.paragraph_format.space_after = Pt(2)
    para.add_run(text)


def add_answer(doc, question, parts):
    h2(doc, question)
    for label, text in parts:
        p(doc, f"{label}: {text}", bold_label=f"{label}:")


def table(doc, headers, rows):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, header in enumerate(headers):
        tbl.rows[0].cells[i].text = header
        for run in tbl.rows[0].cells[i].paragraphs[0].runs:
            run.bold = True
    for row in rows:
        cells = tbl.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    doc.add_paragraph()


ANSWERS = [
    ("Câu 1. Trình bày bài toán tìm kiếm trong AI", [
        ("Khái niệm", "Trong trí tuệ nhân tạo, nhiều bài toán được xem như bài toán tìm kiếm trong không gian trạng thái. Tác tử bắt đầu từ một trạng thái ban đầu, thực hiện các hành động hợp lệ để tạo trạng thái mới và cố gắng đạt tới trạng thái đích. Cách nhìn này xuất hiện trong bài toán tìm đường, 8-puzzle, lập kế hoạch, robot, trò chơi và chứng minh định lý."),
        ("Thành phần", "Một bài toán tìm kiếm gồm initial state, goal test hoặc goal state, tập actions, transition model và path cost. Initial state cho biết điểm xuất phát; actions cho biết các hành động có thể làm; transition model cho biết kết quả sau hành động; goal test kiểm tra đã đạt mục tiêu chưa; path cost đo chi phí của lời giải."),
        ("Nguyên lý giải", "Thuật toán tìm kiếm tạo cây tìm kiếm từ trạng thái ban đầu. Mỗi nút thường chứa state, parent, action, path cost g(n) và depth. Frontier là tập nút đã sinh nhưng chưa mở rộng; explored là tập trạng thái đã xét để tránh lặp. Tùy chiến lược chọn nút trong frontier, ta có BFS, DFS, UCS, Greedy hoặc A*."),
        ("Ví dụ", "Trong bài toán tìm đường từ A đến G, trạng thái là thành phố hiện tại, hành động là đi sang thành phố kề, transition model cho biết thành phố mới sau khi đi, goal state là G và path cost là tổng số km."),
        ("Kết luận", "Bài toán tìm kiếm là nền tảng vì nó biến một vấn đề trừu tượng thành mô hình rõ ràng gồm trạng thái, hành động và mục tiêu, từ đó có thể dùng thuật toán để tìm lời giải một cách hệ thống."),
    ]),
    ("Câu 2. Trình bày BFS", [
        ("Khái niệm", "Breadth-First Search là tìm kiếm theo chiều rộng. Thuật toán mở rộng các nút theo từng mức độ sâu: trước hết xét trạng thái ban đầu, sau đó tất cả trạng thái ở độ sâu 1, rồi độ sâu 2, cứ tiếp tục cho đến khi gặp đích."),
        ("Cấu trúc dữ liệu", "BFS dùng hàng đợi FIFO. Nút sinh trước sẽ được mở rộng trước, nhờ vậy thuật toán duyệt đều theo từng lớp của cây tìm kiếm."),
        ("Tính chất", "BFS hoàn chỉnh nếu branching factor b hữu hạn, nghĩa là nếu tồn tại lời giải thì BFS sẽ tìm thấy. BFS tối ưu khi mọi hành động có chi phí bằng nhau vì lời giải đầu tiên tìm được là lời giải có số bước ít nhất."),
        ("Độ phức tạp", "Nếu lời giải nông nhất ở độ sâu d và mỗi nút trung bình có b nhánh, thời gian và bộ nhớ của BFS thường là O(b^d). Điểm yếu lớn nhất của BFS là bộ nhớ, vì frontier ở mức sâu có thể rất lớn."),
        ("Ứng dụng", "BFS phù hợp cho bài toán cần lời giải ít bước nhất, ví dụ tìm đường trong mê cung không trọng số, duyệt mạng xã hội theo số cạnh hoặc tìm trạng thái gần nhất trong game đơn giản."),
    ]),
    ("Câu 3. Trình bày DFS", [
        ("Khái niệm", "Depth-First Search là tìm kiếm theo chiều sâu. Thuật toán chọn một nhánh và đi sâu nhất có thể, khi không thể đi tiếp hoặc gặp thất bại thì quay lui để thử nhánh khác."),
        ("Cấu trúc dữ liệu", "DFS có thể dùng stack LIFO hoặc đệ quy. Nút sinh sau thường được mở rộng trước, làm thuật toán nhanh chóng đi xuống sâu trong cây."),
        ("Ưu điểm", "DFS tiết kiệm bộ nhớ hơn BFS vì chỉ cần lưu đường đi hiện tại và một số nhánh chưa xét. Không gian bộ nhớ thường O(bm), với m là độ sâu tối đa."),
        ("Nhược điểm", "DFS không tối ưu vì lời giải đầu tiên có thể không phải lời giải ngắn nhất hoặc rẻ nhất. Trong không gian vô hạn hoặc có chu trình, DFS có thể đi mãi nếu không kiểm soát lặp hoặc giới hạn độ sâu."),
        ("Kết luận", "DFS hữu ích khi bộ nhớ hạn chế, không gian tìm kiếm có giới hạn hoặc lời giải có thể nằm sâu, nhưng không phù hợp khi yêu cầu chắc chắn lời giải tối ưu."),
    ]),
    ("Câu 4. Trình bày UCS", [
        ("Khái niệm", "Uniform Cost Search là thuật toán luôn mở rộng nút có chi phí đường đi g(n) nhỏ nhất tính từ trạng thái ban đầu đến nút đó. Nó là lựa chọn tự nhiên khi các hành động có chi phí khác nhau."),
        ("Nguyên lý", "Frontier được tổ chức bằng hàng đợi ưu tiên theo g(n). Mỗi lần lấy ra nút rẻ nhất để mở rộng. Nếu gặp trạng thái đích khi nút đó được lấy ra khỏi hàng đợi ưu tiên, lời giải là tối ưu."),
        ("Tính chất", "UCS hoàn chỉnh và tối ưu nếu chi phí mỗi bước dương, vì thuật toán xét các đường đi theo thứ tự tăng dần của tổng chi phí."),
        ("So sánh", "BFS tối ưu theo số bước khi chi phí bằng nhau; UCS tối ưu theo tổng chi phí ngay cả khi mỗi cạnh có trọng số khác nhau. Tuy nhiên UCS có thể xét nhiều đường chi phí thấp nhưng chưa gần đích."),
        ("Ứng dụng", "UCS thường dùng trong tìm đường có trọng số, định tuyến, lập kế hoạch khi mỗi hành động có chi phí, thời gian hoặc rủi ro khác nhau."),
    ]),
    ("Câu 5. Trình bày Greedy Best-First Search", [
        ("Khái niệm", "Greedy Best-First Search là tìm kiếm có thông tin, dùng heuristic h(n) để ước lượng mức độ gần đích của nút n. Thuật toán luôn chọn nút có h(n) nhỏ nhất."),
        ("Nguyên lý", "Khác với UCS chỉ quan tâm chi phí đã đi g(n), Greedy chỉ quan tâm chi phí ước lượng còn lại h(n). Vì vậy nó có xu hướng lao nhanh về phía được cho là gần đích nhất."),
        ("Ưu điểm", "Nếu heuristic tốt, Greedy có thể tìm lời giải nhanh và mở rộng ít nút hơn tìm kiếm mù. Nó dễ hiểu và dễ triển khai bằng hàng đợi ưu tiên."),
        ("Nhược điểm", "Greedy không đảm bảo tối ưu vì bỏ qua chi phí đã đi. Nó cũng có thể bị dẫn sai nếu heuristic đánh giá sai, thậm chí mắc kẹt trong các vùng có vẻ gần đích nhưng thực chất không tốt."),
        ("Ví dụ", "Trong tìm đường, nếu h(n) là khoảng cách đường chim bay đến đích, Greedy luôn chọn thành phố có khoảng cách đường chim bay nhỏ nhất, dù đường thực tế có thể vòng xa hoặc chi phí cao."),
    ]),
    ("Câu 6. Trình bày A* Search", [
        ("Khái niệm", "A* là thuật toán tìm kiếm có thông tin rất quan trọng. Nó chọn nút theo hàm đánh giá f(n)=g(n)+h(n), trong đó g(n) là chi phí đã đi từ đầu đến n, còn h(n) là chi phí ước lượng từ n đến đích."),
        ("Ý nghĩa công thức", "A* cân bằng giữa quá khứ và tương lai: g(n) giúp không bỏ qua chi phí thực tế đã trả, h(n) giúp hướng tìm kiếm về phía đích. Vì vậy A* thường hiệu quả hơn UCS và đáng tin hơn Greedy."),
        ("Điều kiện tối ưu", "A* tree search tối ưu nếu heuristic admissible, tức h(n) không bao giờ lớn hơn chi phí thật h*(n). Với graph search, heuristic consistent, nghĩa là h(n) <= c(n,a,n') + h(n'), giúp đảm bảo tối ưu và tránh phải mở lại nút."),
        ("Ưu nhược điểm", "Ưu điểm của A* là vừa có tính định hướng vừa có thể tối ưu. Nhược điểm là tốn bộ nhớ vì phải lưu frontier và explored; hiệu quả phụ thuộc rất mạnh vào chất lượng heuristic."),
        ("Ứng dụng", "A* được dùng rộng rãi trong tìm đường GPS, game, robot, lập kế hoạch hành động và các bài toán tối ưu đường đi."),
    ]),
    ("Câu 7. Trình bày Hypothesis Space trong học khái niệm", [
        ("Khái niệm", "Hypothesis Space H là tập tất cả giả thuyết mà thuật toán học được phép xem xét. Trong học khái niệm, mỗi giả thuyết là một mô tả có thể phân loại ví dụ là dương hoặc âm."),
        ("Vai trò", "Kích thước và dạng của H quyết định thuật toán có thể học được gì. Nếu H quá hẹp, giả thuyết đúng có thể không nằm trong H. Nếu H quá rộng, có thể có quá nhiều giả thuyết khớp dữ liệu huấn luyện nhưng tổng quát khác nhau."),
        ("Ví dụ", "Trong ví dụ EnjoySport của Mitchell, mỗi giả thuyết có thể là một bộ điều kiện trên các thuộc tính như Sky, AirTemp, Humidity, Wind, Water và Forecast. Ký hiệu '?' nghĩa là chấp nhận mọi giá trị, còn '0' nghĩa là không chấp nhận giá trị nào."),
        ("Liên hệ", "Thuật toán học tìm một giả thuyết h thuộc H sao cho h phù hợp dữ liệu huấn luyện. Version space là phần còn lại của H sau khi loại bỏ các giả thuyết không nhất quán với dữ liệu."),
        ("Kết luận", "Hypothesis space là khung tìm kiếm của học máy: học thực chất là tìm kiếm trong không gian giả thuyết dưới ràng buộc của dữ liệu."),
    ]),
    ("Câu 8. Trình bày Version Space", [
        ("Khái niệm", "Version Space là tập tất cả giả thuyết trong H còn nhất quán với toàn bộ ví dụ huấn luyện đã quan sát. Một giả thuyết nhất quán nếu nó phân loại đúng mọi ví dụ dương và âm trong tập huấn luyện."),
        ("Ý nghĩa", "Ban đầu version space có thể rất lớn. Khi có thêm ví dụ huấn luyện, các giả thuyết sai bị loại bỏ và version space thu hẹp dần. Nếu dữ liệu đủ tốt và không nhiễu, version space có thể hội tụ về khái niệm đúng."),
        ("Biểu diễn", "Trong Candidate Elimination, version space được biểu diễn bằng hai biên S và G. S là tập giả thuyết cụ thể nhất còn đúng; G là tập giả thuyết tổng quát nhất còn đúng."),
        ("Ưu điểm", "Version space cho ta biết không chỉ một mô hình học được, mà toàn bộ các giả thuyết còn có thể đúng. Điều này giúp hiểu mức độ chắc chắn của quá trình học."),
        ("Nhược điểm", "Nếu dữ liệu nhiễu hoặc gán nhãn sai, version space có thể rỗng. Ngoài ra, với H lớn, việc duy trì toàn bộ version space có thể rất tốn kém."),
    ]),
    ("Câu 9. Trình bày Candidate Elimination", [
        ("Khái niệm", "Candidate Elimination là thuật toán học khái niệm được Tom Mitchell trình bày để duy trì version space bằng hai biên S và G. Thuật toán cập nhật hai biên này mỗi khi nhận ví dụ mới."),
        ("Biên S và G", "S chứa các giả thuyết cụ thể nhất còn nhất quán với dữ liệu. G chứa các giả thuyết tổng quát nhất còn nhất quán. Mọi giả thuyết hợp lệ trong version space nằm giữa S và G theo quan hệ tổng quát hơn hoặc cụ thể hơn."),
        ("Ví dụ dương", "Khi gặp ví dụ dương, mọi giả thuyết trong G không bao phủ ví dụ đó bị loại. Các giả thuyết trong S không bao phủ ví dụ dương sẽ được tổng quát hóa tối thiểu để bao phủ ví dụ, nhưng vẫn không vượt quá G."),
        ("Ví dụ âm", "Khi gặp ví dụ âm, mọi giả thuyết trong S bao phủ ví dụ âm bị loại. Các giả thuyết trong G bao phủ ví dụ âm sẽ được chuyên biệt hóa tối thiểu để loại ví dụ âm, nhưng vẫn tổng quát hơn một giả thuyết nào đó trong S."),
        ("Đánh giá", "Thuật toán minh họa rõ học là quá trình loại bỏ giả thuyết sai. Tuy nhiên nó giả định dữ liệu không nhiễu và khái niệm đích nằm trong H, nên trong thực tế cần thận trọng."),
    ]),
    ("Câu 10. Trình bày biểu diễn tri thức và cơ sở tri thức", [
        ("Khái niệm", "Biểu diễn tri thức là cách mã hóa hiểu biết về thế giới dưới dạng máy tính có thể lưu trữ và suy luận. Một hệ AI không chỉ cần dữ liệu, mà còn cần tri thức về sự vật, quan hệ, quy luật và ngữ cảnh."),
        ("Cơ sở tri thức", "Knowledge Base là tập hợp các câu biểu diễn fact và rule. Fact là sự kiện cụ thể như Student(An). Rule là luật nếu-thì như Student(x) -> Human(x)."),
        ("Suy diễn", "Inference là quá trình rút ra kết luận mới từ KB. Nếu biết Human(Socrates) và forall x Human(x)->Mortal(x), hệ thống suy ra Mortal(Socrates)."),
        ("Yêu cầu", "Một hệ biểu diễn tri thức tốt cần có khả năng biểu đạt đủ mạnh, cú pháp và ngữ nghĩa rõ ràng, hỗ trợ suy diễn đúng đắn, đồng thời đủ hiệu quả để xử lý trong thực tế."),
        ("Ứng dụng", "Biểu diễn tri thức được dùng trong hệ chuyên gia, chẩn đoán y khoa, trợ lý thông minh, chứng minh định lý, ontology và các hệ thống giải thích."),
    ]),
    ("Câu 11. Trình bày Fact, Rule và Knowledge Base", [
        ("Fact", "Fact là một khẳng định cụ thể được xem là đúng trong miền bài toán. Ví dụ: Fever(An), Human(Socrates), Parent(John,Mary). Fact mô tả dữ kiện ban đầu của hệ thống."),
        ("Rule", "Rule là luật suy diễn, thường có dạng nếu điều kiện thì kết luận. Ví dụ: Fever(x) and Cough(x) -> FluRisk(x). Rule cho phép hệ thống tạo tri thức mới từ tri thức đã có."),
        ("Knowledge Base", "KB gồm tập fact và rule. Nó đóng vai trò bộ nhớ tri thức của hệ AI. Hệ suy diễn đọc KB, áp dụng luật và trả lời truy vấn hoặc đưa ra quyết định."),
        ("Ví dụ", "KB có Fever(An), Cough(An), Fever(x) and Cough(x)->FluRisk(x). Từ đó hệ thống suy ra FluRisk(An)."),
        ("Kết luận", "Fact cung cấp dữ liệu cụ thể, rule cung cấp tri thức tổng quát, còn KB kết hợp cả hai để tạo nền tảng suy luận cho AI."),
    ]),
    ("Câu 12. Trình bày logic mệnh đề", [
        ("Khái niệm", "Logic mệnh đề biểu diễn tri thức bằng các mệnh đề nguyên tử chỉ nhận giá trị đúng hoặc sai. Các mệnh đề được kết hợp bằng các phép toán logic như NOT, AND, OR, IMPLIES, IFF."),
        ("Cú pháp và ngữ nghĩa", "Cú pháp quy định cách viết câu hợp lệ, ví dụ P -> Q. Ngữ nghĩa quy định điều kiện đúng sai. Mệnh đề P -> Q chỉ sai khi P đúng và Q sai; đây là điểm dễ nhầm trong bảng chân trị."),
        ("Suy diễn", "Ta nói KB entails alpha nếu mọi mô hình làm KB đúng cũng làm alpha đúng. Có thể kiểm tra bằng bảng chân trị, luật suy diễn hoặc resolution."),
        ("Ưu điểm", "Logic mệnh đề đơn giản, rõ ràng, dễ cài đặt và phù hợp với miền bài toán có số mệnh đề hữu hạn."),
        ("Nhược điểm", "Nó kém biểu đạt vì không nói được trực tiếp về đối tượng, quan hệ và lượng từ. Muốn biểu diễn 'mọi người đều phải chết', logic mệnh đề phải liệt kê nhiều mệnh đề riêng."),
    ]),
    ("Câu 13. Trình bày logic vị từ bậc nhất", [
        ("Khái niệm", "Logic vị từ bậc nhất mở rộng logic mệnh đề bằng đối tượng, thuộc tính, quan hệ, biến, hàm và lượng từ. Nhờ đó nó biểu diễn tri thức tổng quát và tự nhiên hơn."),
        ("Thành phần", "Constant biểu diễn đối tượng cụ thể như An, Socrates. Variable là biến như x, y. Predicate biểu diễn thuộc tính hoặc quan hệ như Student(x), Loves(x,y). Function biểu diễn ánh xạ như MotherOf(x)."),
        ("Ví dụ", "Câu 'mọi người đều phải chết' có thể viết forall x Human(x)->Mortal(x). Câu 'có một sinh viên thích AI' có thể viết exists x Student(x) and Likes(x,AI)."),
        ("Ưu điểm", "Logic vị từ có khả năng biểu đạt mạnh, tránh phải liệt kê từng trường hợp, biểu diễn được quan hệ phức tạp giữa các đối tượng."),
        ("Nhược điểm", "Suy diễn trong logic vị từ phức tạp hơn logic mệnh đề, cần các kỹ thuật như chuẩn hóa, thay thế biến, unification và resolution."),
    ]),
    ("Câu 14. Trình bày lượng từ phổ quát và tồn tại", [
        ("Universal Quantifier", "Lượng từ phổ quát forall có nghĩa là 'với mọi'. Câu forall x Student(x)->Person(x) nghĩa là mọi x, nếu x là sinh viên thì x là người."),
        ("Existential Quantifier", "Lượng từ tồn tại exists có nghĩa là 'tồn tại ít nhất một'. Câu exists x Student(x) and GoodAtAI(x) nghĩa là có ít nhất một sinh viên giỏi AI."),
        ("Phạm vi", "Khi dùng lượng từ cần chú ý phạm vi của biến và dấu ngoặc. Sai phạm vi có thể làm câu mang nghĩa khác. Ví dụ forall x exists y Loves(x,y) khác với exists y forall x Loves(x,y)."),
        ("Liên hệ phủ định", "not forall x P(x) tương đương exists x not P(x). not exists x P(x) tương đương forall x not P(x). Đây là quy tắc quan trọng khi biến đổi logic."),
        ("Kết luận", "Lượng từ giúp logic vị từ biểu diễn các phát biểu tổng quát và tồn tại, là điểm mạnh lớn so với logic mệnh đề."),
    ]),
    ("Câu 15. Trình bày Forward Chaining", [
        ("Khái niệm", "Forward Chaining là suy diễn tiến, bắt đầu từ các dữ kiện đã biết và liên tục áp dụng luật để sinh ra kết luận mới."),
        ("Nguyên lý", "Nếu KB có A và luật A -> B, hệ thống thêm B vào KB. Sau đó nếu có B và B -> C, hệ thống tiếp tục suy ra C. Quá trình dừng khi đạt mục tiêu hoặc không còn luật áp dụng."),
        ("Đặc điểm", "Forward chaining là data-driven, tức được dẫn dắt bởi dữ liệu. Nó phù hợp khi có nhiều dữ kiện và muốn biết tất cả kết luận có thể suy ra."),
        ("Ví dụ", "Trong chẩn đoán, từ dữ kiện sốt, ho, đau họng cùng các luật y khoa, hệ thống có thể suy ra nguy cơ cúm hoặc cần kiểm tra thêm."),
        ("Đánh giá", "Ưu điểm là tự nhiên, có thể suy ra nhiều kết luận. Nhược điểm là có thể sinh nhiều kết luận không liên quan nếu truy vấn chỉ là một mục tiêu hẹp."),
    ]),
    ("Câu 16. Trình bày Backward Chaining", [
        ("Khái niệm", "Backward Chaining là suy diễn lùi, bắt đầu từ mục tiêu cần chứng minh rồi truy ngược về các điều kiện cần thiết."),
        ("Nguyên lý", "Muốn chứng minh B, hệ thống tìm luật có kết luận B, ví dụ A -> B. Khi đó bài toán chuyển thành chứng minh A. Nếu A là fact trong KB hoặc chứng minh được từ luật khác thì B được chứng minh."),
        ("Đặc điểm", "Backward chaining là goal-driven, phù hợp khi người dùng hỏi một truy vấn cụ thể và hệ thống chỉ cần chứng minh truy vấn đó."),
        ("Ví dụ", "Trong Prolog, để chứng minh Mortal(Socrates), hệ thống tìm luật Human(x)->Mortal(x), hợp nhất x với Socrates, rồi kiểm tra Human(Socrates)."),
        ("So sánh", "Forward đi từ dữ kiện đến kết luận và có thể suy ra nhiều điều; backward đi từ mục tiêu về dữ kiện nên tiết kiệm hơn khi chỉ cần trả lời một câu hỏi cụ thể."),
    ]),
    ("Câu 17. Trình bày Resolution", [
        ("Khái niệm", "Resolution là quy tắc suy diễn dựa trên việc triệt tiêu hai literal đối nhau. Từ (A or B) và (not B or C), ta có thể suy ra (A or C)."),
        ("Chứng minh bằng phản chứng", "Để chứng minh KB entails alpha, ta thêm not alpha vào KB rồi chuyển tất cả câu về dạng CNF. Nếu áp dụng resolution mà suy ra mệnh đề rỗng, tức mâu thuẫn, thì KB entails alpha."),
        ("Vai trò", "Resolution là nền tảng của chứng minh định lý tự động vì nó có dạng cơ học, dễ cài đặt và có tính đầy đủ trong logic mệnh đề."),
        ("Trong logic vị từ", "Resolution cần thêm unification để xử lý biến. Hai literal chỉ triệt tiêu được nếu có phép thế làm chúng khớp nhau."),
        ("Kết luận", "Resolution là phương pháp suy diễn quan trọng trong AI logic, đặc biệt trong hệ chứng minh tự động và các hệ dựa trên luật."),
    ]),
    ("Câu 18. Trình bày Unification", [
        ("Khái niệm", "Unification là quá trình tìm phép thay thế biến để hai biểu thức logic trở nên giống nhau. Phép thay thế tìm được thường gọi là substitution."),
        ("Ví dụ", "P(x) và P(Socrates) có thể hợp nhất bằng phép thế x/Socrates. Loves(x,AI) và Loves(An,y) có thể hợp nhất bằng x/An, y/AI."),
        ("Most General Unifier", "MGU là phép hợp nhất tổng quát nhất, không ràng buộc biến nhiều hơn mức cần thiết. Trong suy diễn logic, dùng MGU giúp giữ lời giải tổng quát."),
        ("Vai trò", "Unification là bước cốt lõi của backward chaining, Prolog và resolution trong logic vị từ. Nó cho phép luật tổng quát áp dụng cho dữ kiện cụ thể."),
        ("Kết luận", "Nếu resolution là quy tắc suy diễn, thì unification là cơ chế khớp mẫu giúp quy tắc đó hoạt động với biến và đối tượng trong logic vị từ."),
    ]),
    ("Câu 19. Trình bày Domain Theory và Explanation-Based Learning", [
        ("Domain Theory", "Domain theory là tri thức nền về miền bài toán, thường được biểu diễn bằng các luật. Nó giải thích vì sao một ví dụ thuộc hoặc không thuộc một khái niệm."),
        ("EBL", "Explanation-Based Learning là phương pháp học từ giải thích. Hệ thống dùng domain theory để giải thích một ví dụ huấn luyện, sau đó tổng quát hóa lời giải thích thành một luật dùng cho các trường hợp tương tự."),
        ("Quy trình", "Các bước thường gồm: nhận ví dụ, chứng minh ví dụ bằng domain theory, xác định các điều kiện thật sự cần thiết trong lời giải thích, rồi tạo quy tắc tổng quát."),
        ("Ưu điểm", "EBL có thể học nhanh từ rất ít ví dụ vì tận dụng tri thức nền. Nó giúp biến quá trình suy luận dài thành luật ngắn hơn để dùng lại."),
        ("Nhược điểm", "EBL phụ thuộc mạnh vào chất lượng domain theory. Nếu tri thức nền sai hoặc thiếu, quy tắc học được có thể sai hoặc không đủ hữu ích."),
    ]),
    ("Câu 20. Trình bày FOIL và Inductive Logic Programming", [
        ("ILP", "Inductive Logic Programming là hướng học máy học các luật logic từ ví dụ và tri thức nền. Khác với nhiều thuật toán học thuộc tính phẳng, ILP có thể học quan hệ giữa đối tượng."),
        ("FOIL", "FOIL là thuật toán học luật dạng Horn clause trong logic vị từ. Nó xây dựng luật bằng cách thêm dần literal vào phần điều kiện để bao phủ ví dụ dương và loại bỏ ví dụ âm."),
        ("Nguyên lý", "FOIL bắt đầu từ luật tổng quát, sau đó chuyên biệt hóa bằng cách chọn literal làm tăng khả năng phân biệt giữa ví dụ dương và âm. Khi một luật đủ tốt, thuật toán thêm luật đó vào tập giả thuyết."),
        ("Ưu điểm", "FOIL và ILP biểu diễn được tri thức quan hệ, ví dụ quan hệ gia đình, cấu trúc phân tử, quan hệ trong cơ sở dữ liệu."),
        ("Nhược điểm", "Không gian tìm kiếm luật rất lớn, chi phí tính toán cao và kết quả phụ thuộc nhiều vào predicate, ví dụ huấn luyện và tri thức nền."),
    ]),
    ("Câu 21. Trình bày Machine Learning theo Tom Mitchell", [
        ("Định nghĩa", "Theo Tom Mitchell, một chương trình máy tính được nói là học từ kinh nghiệm E đối với lớp nhiệm vụ T và độ đo hiệu năng P nếu hiệu năng của nó trên T, đo bởi P, được cải thiện nhờ E."),
        ("Phân tích T", "T là task, tức nhiệm vụ hệ thống cần thực hiện, ví dụ phân loại email, dự đoán giá nhà, nhận dạng chữ viết tay hoặc chơi cờ."),
        ("Phân tích P", "P là performance measure, tức thước đo đánh giá mức độ làm tốt nhiệm vụ, ví dụ accuracy, F1-score, mean squared error, lợi nhuận hoặc tỷ lệ thắng."),
        ("Phân tích E", "E là experience, tức kinh nghiệm dùng để học, thường là dữ liệu huấn luyện, ván chơi đã trải qua, phản hồi người dùng hoặc lịch sử quan sát."),
        ("Ví dụ", "Với lọc spam: T là phân loại email spam/không spam, P là độ chính xác hoặc F1, E là tập email đã gán nhãn. Nếu sau khi học từ E, hệ thống phân loại email mới tốt hơn, ta nói nó đã học."),
    ]),
    ("Câu 22. So sánh Supervised, Unsupervised và Reinforcement Learning", [
        ("Supervised Learning", "Học có giám sát dùng dữ liệu có nhãn, mỗi mẫu thường là cặp (x,y). Mục tiêu là học hàm ánh xạ từ đầu vào x sang đầu ra y. Hai dạng chính là regression và classification."),
        ("Unsupervised Learning", "Học không giám sát dùng dữ liệu không nhãn. Mục tiêu là khám phá cấu trúc ẩn trong dữ liệu như phân cụm, giảm chiều, phát hiện bất thường hoặc tìm biểu diễn."),
        ("Reinforcement Learning", "Học tăng cường gồm agent tương tác với environment, chọn action, nhận reward và học policy để tối đa hóa tổng phần thưởng dài hạn."),
        ("Ví dụ", "Supervised: dự đoán giá nhà từ diện tích. Unsupervised: phân nhóm khách hàng theo hành vi mua sắm. Reinforcement: robot học di chuyển hoặc chương trình chơi game."),
        ("Kết luận", "Khác biệt cốt lõi nằm ở tín hiệu học: supervised học từ đáp án đúng, unsupervised học từ cấu trúc dữ liệu, reinforcement học từ thưởng phạt qua tương tác."),
    ]),
    ("Câu 23. Trình bày Linear Regression", [
        ("Khái niệm", "Linear Regression là thuật toán học có giám sát dùng để dự đoán biến liên tục. Mô hình giả định y có quan hệ xấp xỉ tuyến tính với các đặc trưng đầu vào."),
        ("Mô hình", "Dạng tổng quát là h_theta(x)=theta_0+theta_1x_1+...+theta_nx_n, hoặc h_theta(x)=theta^T x nếu thêm x_0=1."),
        ("Hàm mất mát", "Thường dùng mean squared error hoặc tổng bình phương sai số: J(theta)=1/(2m) sum_i (h_theta(x_i)-y_i)^2. Huấn luyện nhằm tìm theta làm J nhỏ nhất."),
        ("Ưu nhược điểm", "Ưu điểm là đơn giản, nhanh, dễ giải thích và là nền tảng của nhiều mô hình thống kê. Nhược điểm là giả định tuyến tính, nhạy với outlier và khó mô hình hóa quan hệ phi tuyến mạnh."),
        ("Ứng dụng", "Dự đoán giá nhà, doanh thu, chi phí, điểm số hoặc các đại lượng liên tục khác."),
    ]),
    ("Câu 24. Trình bày Logistic Regression", [
        ("Khái niệm", "Logistic Regression là mô hình phân loại, thường dùng cho bài toán nhị phân. Dù có tên regression, đầu ra của nó là xác suất thuộc lớp dương."),
        ("Mô hình", "Mô hình tính z=theta^T x rồi đưa qua hàm sigmoid g(z)=1/(1+e^(-z)). Giá trị g(z) nằm trong khoảng 0 đến 1 và được hiểu là P(y=1|x)."),
        ("Ra quyết định", "Nếu xác suất lớn hơn hoặc bằng ngưỡng, thường là 0.5, mô hình dự đoán lớp 1; ngược lại dự đoán lớp 0."),
        ("Ưu nhược điểm", "Ưu điểm là đơn giản, dễ giải thích theo xác suất, hiệu quả khi ranh giới gần tuyến tính. Nhược điểm là khó biểu diễn ranh giới phi tuyến nếu không mở rộng đặc trưng."),
        ("Ứng dụng", "Dự đoán email spam, bệnh/không bệnh, khách hàng rời bỏ, giao dịch gian lận hoặc không gian lận."),
    ]),
    ("Câu 25. Trình bày Decision Tree, Entropy và Information Gain", [
        ("Decision Tree", "Cây quyết định là mô hình dự đoán gồm nút kiểm tra thuộc tính, nhánh biểu diễn kết quả kiểm tra và lá biểu diễn nhãn hoặc giá trị dự đoán."),
        ("Cách học", "Thuật toán xây cây bằng cách chia dữ liệu đệ quy. Tại mỗi nút, chọn thuộc tính giúp các tập con sau khi chia tinh khiết hơn."),
        ("Entropy", "Entropy đo độ hỗn loạn của tập dữ liệu. Với hai lớp, Entropy(S)=-p+log2(p+)-p-log2(p-). Entropy bằng 0 khi tập hoàn toàn cùng lớp và cao khi các lớp trộn lẫn."),
        ("Information Gain", "Information Gain đo mức giảm entropy sau khi chia theo thuộc tính A: Gain(S,A)=Entropy(S)-sum_v |S_v|/|S| Entropy(S_v). Thuộc tính có gain cao thường được chọn trước trong ID3."),
        ("Đánh giá", "Cây quyết định dễ hiểu và dễ giải thích nhưng dễ overfitting nếu cây quá sâu, nên cần pruning, giới hạn độ sâu hoặc dùng ensemble."),
    ]),
    ("Câu 26. Trình bày KNN", [
        ("Khái niệm", "K-Nearest Neighbors là thuật toán dự đoán dựa trên k điểm dữ liệu gần nhất trong tập huấn luyện. Nó thuộc nhóm instance-based learning vì không học mô hình tham số rõ ràng trước."),
        ("Cách hoạt động", "Để dự đoán điểm mới, thuật toán tính khoảng cách từ điểm đó đến các mẫu train, chọn k mẫu gần nhất, rồi bỏ phiếu đa số nếu phân loại hoặc lấy trung bình nếu hồi quy."),
        ("Ảnh hưởng của K", "K nhỏ làm mô hình linh hoạt, bias thấp nhưng variance cao, dễ bị nhiễu. K lớn làm mô hình mượt hơn, variance thấp hơn nhưng bias cao hơn và có thể underfit."),
        ("Ưu nhược điểm", "Ưu điểm là đơn giản, trực quan, không cần huấn luyện phức tạp. Nhược điểm là dự đoán chậm khi dữ liệu lớn, nhạy với thang đo đặc trưng và cần chọn khoảng cách phù hợp."),
        ("Ứng dụng", "Phân loại khách hàng, nhận dạng mẫu, gợi ý đơn giản, phân loại ảnh hoặc văn bản ở mức cơ bản."),
    ]),
    ("Câu 27. Trình bày Naive Bayes và SVM", [
        ("Naive Bayes", "Naive Bayes dựa trên định lý Bayes P(C|x)=P(x|C)P(C)/P(x) và giả định các đặc trưng độc lập có điều kiện theo lớp. Khi phân loại, mô hình chọn lớp có xác suất hậu nghiệm lớn nhất."),
        ("Ưu nhược điểm Naive Bayes", "Naive Bayes rất nhanh, hiệu quả với văn bản và dữ liệu nhiều chiều; nhưng giả định độc lập thường không đúng hoàn toàn và cần smoothing để tránh xác suất bằng 0."),
        ("SVM", "Support Vector Machine tìm siêu phẳng phân tách các lớp sao cho margin lớn nhất. Các điểm gần ranh giới nhất gọi là support vectors vì chúng quyết định vị trí siêu phẳng."),
        ("Kernel", "Kernel trick giúp SVM xử lý ranh giới phi tuyến bằng cách tính độ tương đồng trong không gian đặc trưng cao chiều mà không cần biểu diễn tường minh."),
        ("Ứng dụng", "Naive Bayes hay dùng cho spam và phân loại văn bản. SVM mạnh trong bài toán chiều cao như văn bản, ảnh, dữ liệu sinh học nhưng cần chọn kernel và tham số cẩn thận."),
    ]),
    ("Câu 28. Trình bày Neural Network và Backpropagation", [
        ("Neural Network", "Mạng neural gồm các neuron nhân tạo tổ chức thành input layer, hidden layer và output layer. Mỗi neuron tính tổng có trọng số của đầu vào rồi đưa qua hàm kích hoạt phi tuyến."),
        ("Khả năng học", "Nhờ nhiều lớp và hàm kích hoạt phi tuyến, mạng neural có thể học quan hệ rất phức tạp, đặc biệt trong ảnh, âm thanh, ngôn ngữ và dữ liệu lớn."),
        ("Backpropagation", "Backpropagation tính gradient của hàm mất mát theo từng trọng số bằng quy tắc dây chuyền. Quy trình gồm forward pass, tính loss, backward pass và cập nhật trọng số."),
        ("Cập nhật", "Dạng cập nhật cơ bản là w := w - alpha * dJ/dw, trong đó alpha là learning rate. Learning rate quá lớn dễ không hội tụ, quá nhỏ học chậm."),
        ("Đánh giá", "Neural network mạnh nhưng cần nhiều dữ liệu và tính toán, khó giải thích và dễ overfitting nếu không dùng regularization, dropout, early stopping hoặc dữ liệu đủ lớn."),
    ]),
    ("Câu 29. Trình bày Bias, Variance, Overfitting và Underfitting", [
        ("Bias", "Bias là sai số do giả định mô hình quá đơn giản hoặc lệch so với quy luật thật. Bias cao thường dẫn đến underfitting, tức mô hình không học được cấu trúc quan trọng."),
        ("Variance", "Variance là độ nhạy của mô hình với thay đổi trong tập huấn luyện. Variance cao thường dẫn đến overfitting, tức mô hình học cả nhiễu của training set."),
        ("Underfitting", "Underfitting có dấu hiệu train error cao và test error cao. Cách khắc phục là tăng độ phức tạp mô hình, thêm đặc trưng, huấn luyện lâu hơn hoặc giảm regularization quá mạnh."),
        ("Overfitting", "Overfitting có dấu hiệu train error thấp nhưng validation/test error cao. Cách khắc phục là thêm dữ liệu, regularization, pruning, dropout, early stopping, cross-validation hoặc giảm độ phức tạp."),
        ("Kết luận", "Bias-variance tradeoff nói rằng mô hình quá đơn giản và quá phức tạp đều không tốt. Mục tiêu là chọn mô hình cân bằng để tổng quát hóa tốt trên dữ liệu mới."),
    ]),
    ("Câu 30. Trình bày Cross Validation, Regularization và Q-Learning", [
        ("Cross Validation", "Cross-validation đánh giá mô hình bằng cách chia dữ liệu thành nhiều phần. Với k-fold, mỗi lần dùng k-1 phần để train và 1 phần để validation, sau đó lấy trung bình kết quả."),
        ("Regularization", "Regularization thêm thành phần phạt độ phức tạp vào loss, ví dụ J_lambda(theta)=J(theta)+lambda R(theta). Lambda càng lớn thì mô hình càng bị hạn chế, giúp giảm overfitting."),
        ("L1 và L2", "L1 có xu hướng tạo nghiệm thưa, hỗ trợ chọn đặc trưng. L2 làm trọng số nhỏ và mô hình mượt hơn. Cả hai đều giúp kiểm soát variance."),
        ("Q-Learning", "Q-Learning là thuật toán reinforcement learning học giá trị Q(s,a). Công thức cập nhật: Q(s,a):=Q(s,a)+alpha[r+gamma max_a' Q(s',a')-Q(s,a)]."),
        ("Kết luận", "Cross-validation giúp chọn mô hình và hyperparameter khách quan; regularization giúp tổng quát hóa; Q-learning minh họa cách agent học hành động tối ưu qua reward."),
    ]),
]


def build():
    if TARGET.exists() and not BACKUP.exists():
        copy2(TARGET, BACKUP)

    doc = setup_doc()
    title(doc, "ĐỀ CƯƠNG TỰ LUẬN AI - CHƯƠNG 2, 3, 5")
    p(doc, "Bản trả lời chi tiết các câu hỏi trong đề cương, biên soạn theo hướng làm bài tự luận. Nội dung dựa trên các giáo trình/tài liệu: Chon_main_notes, Tom M. Mitchell - Machine Learning, và An Introduction to Statistical Learning.", first_line=False)
    p(doc, "Mỗi câu đều trình bày theo cấu trúc: khái niệm, nguyên lý hoặc công thức, ví dụ minh họa, ưu nhược điểm hoặc ứng dụng, và kết luận.", first_line=False)

    h1(doc, "PHẦN 1. TÓM TẮT Ý CHÍNH CẦN NHỚ")
    h2(doc, "Chương 2: Giải quyết vấn đề bằng tìm kiếm")
    for item in [
        "Bài toán tìm kiếm gồm initial state, actions, transition model, goal test và path cost.",
        "BFS dùng queue FIFO, hoàn chỉnh và tối ưu khi chi phí bằng nhau nhưng tốn bộ nhớ.",
        "DFS dùng stack LIFO hoặc đệ quy, tiết kiệm bộ nhớ nhưng không tối ưu.",
        "UCS chọn g(n) nhỏ nhất; Greedy chọn h(n) nhỏ nhất; A* chọn f(n)=g(n)+h(n).",
        "A* tối ưu khi heuristic admissible; graph search thường cần consistent để ổn định.",
        "Concept learning xem học như tìm kiếm trong hypothesis space; version space là tập giả thuyết còn nhất quán.",
    ]:
        bullet(doc, item)

    h2(doc, "Chương 3: Biểu diễn tri thức và suy diễn")
    for item in [
        "Knowledge base gồm fact và rule; inference dùng tri thức đã có để rút ra tri thức mới.",
        "Logic mệnh đề đơn giản nhưng kém biểu đạt; logic vị từ có đối tượng, quan hệ và lượng từ.",
        "Forward chaining là data-driven; backward chaining là goal-driven.",
        "Resolution thường dùng để chứng minh bằng phản chứng; unification giúp khớp biểu thức có biến.",
        "Domain theory, EBL, FOIL và ILP kết hợp tri thức logic với học máy.",
    ]:
        bullet(doc, item)

    h2(doc, "Chương 5: Học máy")
    for item in [
        "Định nghĩa Tom Mitchell: học từ E đối với T và P nếu hiệu năng P trên T cải thiện nhờ E.",
        "Supervised học từ nhãn; unsupervised tìm cấu trúc; reinforcement học qua reward.",
        "Linear/logistic regression, decision tree, KNN, Naive Bayes, SVM và neural network là các mô hình cần nắm.",
        "Overfitting, underfitting, bias, variance, cross-validation và regularization liên quan trực tiếp tới generalization.",
    ]:
        bullet(doc, item)

    h1(doc, "PHẦN 2. TRẢ LỜI CHI TIẾT CÁC CÂU TỰ LUẬN")
    for question, parts in ANSWERS:
        add_answer(doc, question, parts)

    h1(doc, "PHẦN 3. BẢNG SO SÁNH NHANH")
    table(doc, ["Thuật toán", "Tiêu chí chọn", "Tối ưu?", "Điểm yếu chính"], [
        ("BFS", "Node nông nhất", "Có, nếu chi phí bằng nhau", "Tốn bộ nhớ"),
        ("DFS", "Node sâu nhất", "Không", "Có thể kẹt, không tối ưu"),
        ("UCS", "g(n) nhỏ nhất", "Có, nếu chi phí dương", "Có thể chậm"),
        ("Greedy", "h(n) nhỏ nhất", "Không", "Phụ thuộc heuristic"),
        ("A*", "g(n)+h(n) nhỏ nhất", "Có, nếu heuristic phù hợp", "Tốn bộ nhớ"),
    ])
    table(doc, ["Vấn đề ML", "Dấu hiệu", "Cách khắc phục"], [
        ("Underfitting", "Train error cao, test error cao", "Tăng độ phức tạp, thêm đặc trưng, giảm regularization"),
        ("Overfitting", "Train error thấp, test error cao", "Regularization, thêm dữ liệu, pruning, cross-validation"),
        ("Bias cao", "Mô hình quá đơn giản", "Dùng mô hình linh hoạt hơn"),
        ("Variance cao", "Quá nhạy với dữ liệu train", "Giảm độ phức tạp, regularization, thêm dữ liệu"),
    ])

    h1(doc, "TÀI LIỆU THAM KHẢO")
    for item in [
        "Chon_main_notes.pdf - tài liệu ghi chú/bài giảng chính về AI.",
        "Tom M. Mitchell, Machine Learning, McGraw-Hill, 1997 - định nghĩa T, P, E; concept learning; candidate elimination; decision tree; neural networks; Bayesian learning.",
        "Gareth James, Daniela Witten, Trevor Hastie, Robert Tibshirani, An Introduction to Statistical Learning with Applications in R - regression, classification, KNN, model evaluation, overfitting, bias-variance, regularization.",
    ]:
        bullet(doc, item)

    doc.save(OUTPUT)

    try:
        doc.save(TARGET)
        print(f"UPDATED_TARGET={TARGET}")
    except PermissionError:
        print(f"TARGET_LOCKED={TARGET}")
    except OSError as exc:
        print(f"TARGET_WRITE_FAILED={exc}")

    print(f"OUTPUT={OUTPUT}")
    if BACKUP.exists():
        print(f"BACKUP={BACKUP}")
    print(f"ANSWERS={len(ANSWERS)}")


if __name__ == "__main__":
    build()
