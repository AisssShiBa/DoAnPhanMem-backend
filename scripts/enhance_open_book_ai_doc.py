from pathlib import Path
from shutil import copy2

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


TARGET = Path(r"D:\GK_AI\CuoiKiAI.docx")
BACKUP = TARGET.with_name("CuoiKiAI_backup_truoc_khi_bo_sung_de_mo.docx")


def normalize_font(doc: Document):
    for style_name in ["Normal", "Heading 1", "Heading 2", "Heading 3"]:
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        if style_name == "Normal":
            style.font.size = Pt(12)


def h1(doc, text):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(10)
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(14)


def h2(doc, text):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(8)
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(13)


def h3(doc, text):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(6)
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(12)


def p(doc, text, label=None, indent=True):
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(4)
    para.paragraph_format.line_spacing = 1.08
    if indent:
        para.paragraph_format.first_line_indent = Inches(0.25)
    if label and text.startswith(label):
        run = para.add_run(label)
        run.bold = True
        para.add_run(text[len(label):])
    else:
        para.add_run(text)


def bullet(doc, text):
    para = doc.add_paragraph(style="List Bullet")
    para.paragraph_format.space_after = Pt(2)
    para.add_run(text)


def table(doc, headers, rows):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, header in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        cell.text = header
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for row in rows:
        cells = tbl.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    doc.add_paragraph()


def already_added(doc):
    return any("PHẦN 5. BỔ SUNG CHO ĐỀ MỞ" in para.text for para in doc.paragraphs)


OPEN_BOOK_QA = [
    ("Câu mở rộng 1. So sánh Tree Search và Graph Search",
     "Tree Search xem mỗi đường đi sinh ra là một nhánh riêng trong cây, nên cùng một trạng thái có thể xuất hiện nhiều lần qua nhiều đường khác nhau. Graph Search thêm explored/closed list để ghi nhớ trạng thái đã xét, nhờ đó tránh lặp và giảm mở rộng dư thừa. Khi không gian có chu trình, graph search thường an toàn hơn. Tuy nhiên với A*, graph search muốn tối ưu cần heuristic consistent hoặc có cơ chế mở lại nút khi tìm được đường rẻ hơn."),
    ("Câu mở rộng 2. So sánh Depth-Limited Search và Iterative Deepening Search",
     "Depth-Limited Search là DFS có giới hạn độ sâu l, giúp tránh đi vô hạn nhưng có thể bỏ lỡ lời giải nếu l nhỏ hơn độ sâu lời giải. Iterative Deepening Search chạy DLS nhiều lần với giới hạn tăng dần 0,1,2,... IDS kết hợp ưu điểm BFS và DFS: hoàn chỉnh, tối ưu khi chi phí bằng nhau và dùng ít bộ nhớ. Điểm trừ là duyệt lại các mức nông nhiều lần, nhưng chi phí lặp thường chấp nhận được khi branching factor lớn."),
    ("Câu mở rộng 3. Khi nào heuristic này tốt hơn heuristic khác?",
     "Một heuristic h2 được xem là dominates h1 nếu với mọi n, h2(n) >= h1(n) và cả hai đều admissible. Khi đó h2 gần chi phí thật hơn mà vẫn không vượt quá chi phí thật, nên A* với h2 thường mở rộng ít nút hơn A* với h1. Trong bài thi có thể kết luận: heuristic càng thông tin nhiều, càng sát chi phí thật và vẫn admissible thì càng tốt."),
    ("Câu mở rộng 4. Phân biệt admissible và consistent bằng lời dễ nhớ",
     "Admissible là không nói quá: h(n) không được lớn hơn chi phí thật từ n đến đích. Consistent là không nhảy vọt: ước lượng tại n không được lớn hơn chi phí đi sang n' cộng ước lượng tại n'. Consistent mạnh hơn admissible trong graph search. Nếu heuristic consistent thì f(n) dọc theo đường đi không giảm, giúp A* xử lý ổn định hơn."),
    ("Câu mở rộng 5. Tại sao Candidate Elimination nhạy với nhiễu?",
     "Candidate Elimination giữ lại chỉ những giả thuyết hoàn toàn nhất quán với dữ liệu. Nếu có một ví dụ bị gán nhãn sai, thuật toán có thể loại bỏ giả thuyết đúng khỏi version space. Vì vậy version space có thể rỗng dù bản chất bài toán vẫn có quy luật. Đây là lý do các phương pháp học thực tế thường chấp nhận sai số thay vì yêu cầu khớp tuyệt đối mọi mẫu."),
    ("Câu mở rộng 6. So sánh Find-S và Candidate Elimination",
     "Find-S chỉ học giả thuyết cụ thể nhất phù hợp các ví dụ dương và bỏ qua ví dụ âm, nên đơn giản nhưng không cho biết toàn bộ tập giả thuyết có thể đúng. Candidate Elimination dùng cả ví dụ dương và âm để cập nhật hai biên S và G, biểu diễn toàn bộ version space. Candidate Elimination đầy đủ hơn nhưng phức tạp và nhạy với nhiễu hơn."),
    ("Câu mở rộng 7. So sánh học quy nạp và suy diễn logic",
     "Suy diễn đi từ luật tổng quát và dữ kiện cụ thể để rút ra kết luận chắc chắn nếu luật đúng. Học quy nạp đi từ dữ liệu quan sát để tạo ra luật hoặc mô hình tổng quát, nên kết luận có tính xác suất và có thể sai trên dữ liệu mới. Machine learning chủ yếu là quy nạp; logic inference chủ yếu là suy diễn."),
    ("Câu mở rộng 8. Phân biệt syntax, semantics, entailment, inference",
     "Syntax là quy tắc viết câu hợp lệ. Semantics là ý nghĩa và điều kiện đúng sai của câu. Entailment KB |= alpha là quan hệ ngữ nghĩa: trong mọi mô hình làm KB đúng thì alpha đúng. Inference KB |- alpha là quá trình thuật toán suy ra alpha. Thuật toán sound nếu suy ra gì cũng đúng theo entailment; complete nếu mọi điều được entailment thì thuật toán suy ra được."),
    ("Câu mở rộng 9. Khi nào dùng Forward Chaining, khi nào dùng Backward Chaining?",
     "Dùng forward chaining khi dữ kiện mới liên tục xuất hiện và ta muốn suy ra mọi kết luận có thể, ví dụ hệ cảnh báo. Dùng backward chaining khi có truy vấn cụ thể cần chứng minh, ví dụ hệ hỏi đáp hoặc Prolog. Trong bài thi nên nêu forward là data-driven, backward là goal-driven."),
    ("Câu mở rộng 10. Vì sao logic vị từ mạnh hơn logic mệnh đề?",
     "Logic mệnh đề xem mỗi câu nguyên tử là một khối đúng/sai, không biểu diễn được cấu trúc bên trong. Logic vị từ có đối tượng, thuộc tính, quan hệ, biến và lượng từ, nên biểu diễn được phát biểu tổng quát như forall x Human(x)->Mortal(x). Vì mạnh hơn nên suy diễn trong logic vị từ cũng phức tạp hơn."),
    ("Câu mở rộng 11. So sánh generative và discriminative model",
     "Mô hình generative học P(x,y) hoặc P(x|y)P(y), tức mô hình hóa cách dữ liệu được sinh ra; Naive Bayes là ví dụ. Mô hình discriminative học trực tiếp P(y|x) hoặc ranh giới quyết định giữa các lớp; Logistic Regression và SVM là ví dụ. Generative thường hữu ích khi cần mô hình hóa phân phối dữ liệu; discriminative thường mạnh cho phân loại khi có đủ dữ liệu."),
    ("Câu mở rộng 12. So sánh parametric và non-parametric model",
     "Parametric model có số tham số cố định sau khi chọn mô hình, ví dụ linear regression, logistic regression. Non-parametric model có độ phức tạp tăng theo dữ liệu hoặc không bị cố định bởi số tham số nhỏ, ví dụ KNN, decision tree. Parametric thường nhanh và dễ giải thích hơn; non-parametric linh hoạt hơn nhưng dễ tốn bộ nhớ/tính toán."),
    ("Câu mở rộng 13. Phân biệt loss function, cost function và objective function",
     "Loss function đo sai số trên một mẫu. Cost function thường là trung bình/tổng loss trên toàn bộ tập huấn luyện. Objective function là hàm mục tiêu cần tối ưu, có thể gồm cost cộng regularization. Ví dụ trong ridge regression, objective = MSE + lambda * tổng bình phương trọng số."),
    ("Câu mở rộng 14. Vì sao cần train, validation và test set?",
     "Training set dùng để học tham số. Validation set dùng chọn mô hình và hyperparameter. Test set chỉ dùng đánh giá cuối cùng khả năng tổng quát. Nếu dùng test set nhiều lần để chọn mô hình, test error không còn khách quan vì mô hình đã gián tiếp được điều chỉnh theo test set."),
    ("Câu mở rộng 15. So sánh classification và regression",
     "Regression dự đoán giá trị liên tục như giá nhà hoặc nhiệt độ; metric thường là MSE, RMSE, MAE, R^2. Classification dự đoán nhãn rời rạc như spam/không spam; metric thường là accuracy, precision, recall, F1, AUC. Cả hai đều thuộc supervised learning nếu dữ liệu có nhãn."),
    ("Câu mở rộng 16. Khi nào accuracy gây hiểu nhầm?",
     "Accuracy gây hiểu nhầm khi dữ liệu mất cân bằng. Ví dụ 99% giao dịch là bình thường, mô hình luôn dự đoán bình thường vẫn đạt 99% accuracy nhưng không phát hiện gian lận nào. Khi đó cần dùng precision, recall, F1 hoặc AUC, tùy mục tiêu giảm false positive hay false negative."),
    ("Câu mở rộng 17. So sánh L1 và L2 Regularization",
     "L1 phạt tổng trị tuyệt đối trọng số, có xu hướng đưa nhiều trọng số về 0 nên hỗ trợ chọn đặc trưng. L2 phạt tổng bình phương trọng số, làm trọng số nhỏ và ổn định hơn nhưng ít tạo đúng 0. Cả hai đều giảm overfitting bằng cách hạn chế độ phức tạp mô hình."),
    ("Câu mở rộng 18. Pruning trong Decision Tree là gì?",
     "Pruning là cắt tỉa cây để giảm overfitting. Pre-pruning dừng sớm bằng giới hạn độ sâu, số mẫu tối thiểu hoặc gain tối thiểu. Post-pruning xây cây lớn rồi cắt các nhánh không cải thiện validation performance. Ý tưởng là cây quá sâu học cả nhiễu của training set."),
    ("Câu mở rộng 19. So sánh Bagging và Boosting",
     "Bagging huấn luyện nhiều mô hình độc lập trên các mẫu bootstrap rồi lấy trung bình/bỏ phiếu, giúp giảm variance; Random Forest là ví dụ. Boosting huấn luyện tuần tự, mô hình sau tập trung vào lỗi của mô hình trước, giúp giảm bias nhưng dễ nhạy với nhiễu hơn; AdaBoost/Gradient Boosting là ví dụ."),
    ("Câu mở rộng 20. K-means hoạt động thế nào và hạn chế gì?",
     "K-means chọn K tâm cụm, gán mỗi điểm vào tâm gần nhất, cập nhật tâm bằng trung bình các điểm trong cụm, rồi lặp đến khi ổn định. Hạn chế: phải chọn K trước, nhạy với khởi tạo và outlier, không tốt với cụm không lồi hoặc mật độ khác nhau. Cần chuẩn hóa dữ liệu nếu thang đo đặc trưng khác nhau."),
    ("Câu mở rộng 21. PCA là gì và dùng khi nào?",
     "PCA là phương pháp giảm chiều không giám sát, tìm các hướng phương sai lớn nhất của dữ liệu rồi chiếu dữ liệu lên các hướng đó. PCA giúp nén dữ liệu, trực quan hóa, giảm nhiễu và giảm chi phí tính toán. Nhược điểm là thành phần chính là tổ hợp tuyến tính nên có thể khó giải thích."),
    ("Câu mở rộng 22. Gradient Descent cần chú ý gì?",
     "Gradient descent cập nhật tham số theo hướng giảm hàm mất mát: theta := theta - alpha * gradient J(theta). Learning rate quá lớn có thể dao động hoặc phân kỳ; quá nhỏ hội tụ chậm. Cần chuẩn hóa đặc trưng để tối ưu ổn định hơn, đặc biệt trong linear/logistic regression và neural networks."),
    ("Câu mở rộng 23. So sánh batch, stochastic và mini-batch gradient descent",
     "Batch gradient descent dùng toàn bộ dữ liệu mỗi lần cập nhật, ổn định nhưng chậm khi dữ liệu lớn. Stochastic gradient descent dùng một mẫu mỗi lần, nhanh và nhiễu hơn. Mini-batch dùng một nhóm nhỏ, là lựa chọn phổ biến vì cân bằng tốc độ, ổn định và tận dụng tính toán vector/GPU."),
    ("Câu mở rộng 24. Reinforcement Learning khác supervised learning thế nào?",
     "Supervised learning có nhãn đúng cho từng mẫu, còn reinforcement learning chỉ nhận reward sau hành động, thường bị trễ và phụ thuộc chuỗi quyết định. RL phải cân bằng exploration và exploitation, đồng thời tối ưu tổng reward dài hạn chứ không chỉ dự đoán đúng từng mẫu độc lập."),
    ("Câu mở rộng 25. Q-learning là off-policy nghĩa là gì?",
     "Q-learning học giá trị tối ưu bằng mục tiêu r + gamma max_a' Q(s',a'), tức giả định sau trạng thái kế tiếp sẽ chọn hành động tốt nhất, dù trong quá trình khám phá agent có thể hành động theo epsilon-greedy. Vì học chính sách tối ưu độc lập với chính sách đang dùng để thu thập dữ liệu, Q-learning được gọi là off-policy."),
]


def add_open_book_section(doc):
    doc.add_page_break()
    h1(doc, "PHẦN 5. BỔ SUNG CHO ĐỀ MỞ")
    p(doc, "Đánh giá của mình: file hiện tại đã ổn để đi thi nếu đề hỏi đúng các câu cơ bản. Tuy nhiên vì đây là đề mở được mang tài liệu vào, tài liệu nên đóng vai trò như một bộ tra cứu nhanh: có câu trả lời dài, có bảng chọn thuật toán, có ví dụ chứng minh, có công thức và có các câu hỏi biến dạng. Phần dưới đây bổ sung theo hướng đó.")

    h2(doc, "1. Mục lục tra cứu nhanh theo từ khóa")
    table(doc, ["Nếu đề có từ khóa", "Mở ngay phần", "Ý chính cần viết"], [
        ("state space, tìm đường, puzzle", "Chương 2 - Tìm kiếm", "Initial state, actions, transition model, goal test, path cost"),
        ("queue, ít bước nhất", "BFS", "FIFO, hoàn chỉnh, tối ưu khi chi phí bằng nhau, tốn bộ nhớ"),
        ("stack, quay lui", "DFS", "LIFO/đệ quy, ít bộ nhớ, không tối ưu"),
        ("cost, trọng số", "UCS/A*", "UCS theo g(n), A* theo g(n)+h(n)"),
        ("heuristic", "Greedy/A*", "Admissible, consistent, h(n), f(n)"),
        ("fact, rule, KB", "Biểu diễn tri thức", "Fact + rule + inference"),
        ("forall, exists", "Logic vị từ", "Đối tượng, quan hệ, lượng từ"),
        ("Prolog, truy vấn", "Backward chaining", "Goal-driven, chứng minh mục tiêu"),
        ("T, P, E", "Machine Learning", "Định nghĩa Mitchell"),
        ("train tốt test kém", "Overfitting", "Variance cao, regularization, cross-validation"),
        ("train và test đều kém", "Underfitting", "Bias cao, tăng độ phức tạp"),
        ("spam, Bayes", "Naive Bayes", "Bayes + độc lập có điều kiện"),
        ("margin, kernel", "SVM", "Siêu phẳng, support vector, kernel trick"),
        ("reward, policy", "Reinforcement Learning", "Agent, environment, reward, Q-learning"),
    ])

    h2(doc, "2. Công thức cần có sẵn để chép nhanh")
    table(doc, ["Chủ đề", "Công thức", "Giải thích ký hiệu"], [
        ("A*", "f(n)=g(n)+h(n)", "g(n): chi phí đã đi; h(n): ước lượng đến đích"),
        ("Admissible", "h(n) <= h*(n)", "Heuristic không vượt quá chi phí thật"),
        ("Consistent", "h(n) <= c(n,a,n') + h(n')", "Ước lượng không được nhảy vọt qua một cạnh"),
        ("Entropy", "H(S)=-sum p_i log2(p_i)", "Độ hỗn loạn/không tinh khiết của tập S"),
        ("Information Gain", "Gain(S,A)=H(S)-sum_v |S_v|/|S| H(S_v)", "Mức giảm entropy khi chia theo A"),
        ("Linear Regression", "h_theta(x)=theta^T x", "Dự đoán giá trị liên tục"),
        ("MSE cost", "J(theta)=1/(2m) sum(h_theta(x_i)-y_i)^2", "Sai số bình phương trung bình có hệ số 1/2 để dễ đạo hàm"),
        ("Sigmoid", "g(z)=1/(1+e^(-z))", "Đưa giá trị z về xác suất 0..1"),
        ("Bayes", "P(A|B)=P(B|A)P(A)/P(B)", "Cập nhật xác suất khi biết bằng chứng B"),
        ("Precision", "TP/(TP+FP)", "Trong dự đoán dương, bao nhiêu là đúng"),
        ("Recall", "TP/(TP+FN)", "Trong dương thật, phát hiện được bao nhiêu"),
        ("F1", "2PR/(P+R)", "Trung bình điều hòa của precision và recall"),
        ("Q-learning", "Q:=Q+alpha[r+gamma max Q(s',a')-Q]", "Cập nhật giá trị hành động"),
    ])

    h2(doc, "3. Các câu hỏi mở rộng nên có trong đề mở")
    for q, answer in OPEN_BOOK_QA:
        h3(doc, q)
        p(doc, answer)

    h2(doc, "4. Bộ ví dụ nhỏ có thể chép vào nhiều bài")
    p(doc, "Ví dụ tìm kiếm: Robot đi từ phòng A đến phòng G. Nếu mọi bước đi có chi phí như nhau, BFS tìm đường ít bước nhất. Nếu mỗi hành lang có độ dài khác nhau, UCS tìm đường tổng chi phí nhỏ nhất. Nếu có ước lượng khoảng cách còn lại đến G, A* dùng f(n)=g(n)+h(n) để vừa xét chi phí đã đi vừa xét triển vọng đến đích.", "Ví dụ tìm kiếm:")
    p(doc, "Ví dụ logic: KB có Human(Socrates) và forall x Human(x)->Mortal(x). Bằng backward chaining, để chứng minh Mortal(Socrates), hệ thống tìm luật có kết luận Mortal(x), hợp nhất x với Socrates, rồi kiểm tra Human(Socrates).", "Ví dụ logic:")
    p(doc, "Ví dụ học máy theo Mitchell: Với hệ lọc spam, T là phân loại email, P là F1-score hoặc accuracy, E là tập email đã gán nhãn. Nếu sau khi học từ E, P trên email mới tăng, ta nói chương trình đã học.", "Ví dụ học máy theo Mitchell:")
    p(doc, "Ví dụ overfitting: Một cây quyết định quá sâu phân loại đúng gần như toàn bộ training set nhưng sai nhiều trên test set. Cách xử lý là pruning, giới hạn độ sâu, thêm dữ liệu hoặc dùng cross-validation chọn tham số.", "Ví dụ overfitting:")
    p(doc, "Ví dụ mất cân bằng dữ liệu: Trong phát hiện bệnh hiếm, mô hình luôn dự đoán không bệnh có thể accuracy rất cao nhưng recall bằng 0. Vì vậy cần dùng recall/F1 thay vì chỉ accuracy.", "Ví dụ mất cân bằng dữ liệu:")

    h2(doc, "5. Cách viết câu trả lời dài để ăn điểm trong đề mở")
    bullet(doc, "Bước 1: Định nghĩa khái niệm bằng 2-3 câu rõ ràng.")
    bullet(doc, "Bước 2: Nêu công thức hoặc thuật toán từng bước nếu có.")
    bullet(doc, "Bước 3: Giải thích ý nghĩa từng ký hiệu, tránh chỉ chép công thức.")
    bullet(doc, "Bước 4: So sánh với ít nhất một phương pháp gần giống.")
    bullet(doc, "Bước 5: Đưa ví dụ cụ thể, càng gần đời sống càng dễ được điểm.")
    bullet(doc, "Bước 6: Nêu ưu điểm, nhược điểm và điều kiện áp dụng.")
    bullet(doc, "Bước 7: Kết luận bằng câu chọn phương pháp theo tình huống.")

    h2(doc, "6. Dàn ý siêu dài cho 3 câu tổng hợp có thể ra")
    p(doc, "Câu tổng hợp về tìm kiếm: Trình bày AI giải quyết vấn đề bằng cách mô hình hóa không gian trạng thái. Nêu initial state, actions, transition model, goal test, path cost, frontier, explored. Phân loại uninformed search và informed search. So sánh BFS, DFS, UCS, Greedy, A*. Nêu heuristic, admissible, consistent. Đưa ví dụ tìm đường. Kết luận chọn thuật toán theo chi phí, bộ nhớ và heuristic.", "Câu tổng hợp về tìm kiếm:")
    p(doc, "Câu tổng hợp về tri thức và suy diễn: Mở bài về vai trò của tri thức trong AI. Nêu KB gồm fact/rule, syntax/semantics, entailment/inference. So sánh logic mệnh đề và logic vị từ. Giải thích lượng từ, Horn clause, forward chaining, backward chaining, resolution, unification. Đưa ví dụ Socrates hoặc chẩn đoán bệnh. Kết luận ứng dụng trong hệ chuyên gia và Prolog.", "Câu tổng hợp về tri thức và suy diễn:")
    p(doc, "Câu tổng hợp về học máy: Mở bài bằng định nghĩa Mitchell T-P-E. Phân loại supervised, unsupervised, reinforcement. Nêu quy trình ML: dữ liệu, đặc trưng, train/validation/test, chọn mô hình, loss, tối ưu, đánh giá. Trình bày linear/logistic regression, decision tree, KNN, Naive Bayes, SVM, neural network. Kết thúc bằng overfitting, underfitting, bias-variance, regularization, cross-validation và metric.", "Câu tổng hợp về học máy:")

    h2(doc, "7. Nếu còn thời gian trước khi thi, nên ưu tiên học theo thứ tự")
    bullet(doc, "Ưu tiên 1: A*, BFS/DFS/UCS/Greedy và bảng so sánh tìm kiếm.")
    bullet(doc, "Ưu tiên 2: Forward/Backward chaining, logic mệnh đề/vị từ, resolution/unification.")
    bullet(doc, "Ưu tiên 3: Định nghĩa Mitchell, supervised/unsupervised/RL, overfitting/underfitting.")
    bullet(doc, "Ưu tiên 4: Decision tree, entropy, information gain, KNN, Naive Bayes, SVM.")
    bullet(doc, "Ưu tiên 5: Neural network, backpropagation, regularization, cross-validation, Q-learning.")
    bullet(doc, "Ưu tiên cuối: Các câu mở rộng như PCA, bagging/boosting, generative/discriminative nếu đề có hỏi nâng cao.")

    h2(doc, "8. Nhận xét cuối cùng về mức độ sẵn sàng")
    p(doc, "Với bản đã bổ sung này, tài liệu đủ ổn cho đề mở vì có cả phần trả lời trực tiếp, phần so sánh, phần công thức, phần tình huống và phần câu hỏi mở rộng. Khi vào thi, không nên đọc từ đầu đến cuối; hãy dùng mục lục từ khóa, tìm đúng nhóm câu hỏi, rồi ghép định nghĩa + công thức + ví dụ + so sánh + kết luận.")


def main():
    if not TARGET.exists():
        raise FileNotFoundError(f"Không tìm thấy file: {TARGET}")
    if not BACKUP.exists():
        copy2(TARGET, BACKUP)

    doc = Document(TARGET)
    normalize_font(doc)
    if already_added(doc):
        print("ALREADY_ADDED=1")
        print(f"TARGET={TARGET}")
        print(f"BACKUP={BACKUP}")
        return

    add_open_book_section(doc)
    doc.save(TARGET)
    print(f"UPDATED={TARGET}")
    print(f"BACKUP={BACKUP}")
    print(f"ADDED_QA={len(OPEN_BOOK_QA)}")


if __name__ == "__main__":
    main()
