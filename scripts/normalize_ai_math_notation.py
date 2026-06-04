from pathlib import Path
from shutil import copy2

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt


TARGET = Path(r"D:\GK_AI\CuoiKiAI.docx")
BACKUP = TARGET.with_name("CuoiKiAI_backup_truoc_khi_chuan_hoa_cong_thuc.docx")
OUTPUT_IF_LOCKED = TARGET.with_name("CuoiKiAI_DA_CHUAN_HOA_CONG_THUC.docx")


DIRECT_REPLACEMENTS = {
    "f(n)=g(n)+h(n)": "f(n) = g(n) + h(n)",
    "g(n)+h(n)": "g(n) + h(n)",
    "h(n) <= h*(n)": "h(n) ≤ h*(n)",
    "h(n) <= c(n,a,n') + h(n')": "h(n) ≤ c(n,a,n′) + h(n′)",
    "f(n)<=C*<g(G)": "f(n) ≤ C* < g(G)",
    "f(n)<=C*": "f(n) ≤ C*",
    "g(G)>C*": "g(G) > C*",
    "h2(n) >= h1(n)": "h₂(n) ≥ h₁(n)",
    "h2 dominates h1": "h₂ dominates h₁",
    "h2 được gọi là dominates h1": "h₂ được gọi là dominates h₁",
    "h2 gần chi phí thật hơn h1": "h₂ gần chi phí thật hơn h₁",
    "A* với h2": "A* với h₂",
    "A* với h1": "A* với h₁",
    "h_i >=g h_j": "hᵢ ≥g hⱼ",
    "h_i": "hᵢ",
    "h_j": "hⱼ",
    "PositiveSet(h_j) subset PositiveSet(h_i)": "PositiveSet(hⱼ) ⊆ PositiveSet(hᵢ)",
    "theta_0": "θ₀",
    "theta_1": "θ₁",
    "theta_n": "θₙ",
    "theta^T": "θᵀ",
    "theta": "θ",
    "h_theta": "hθ",
    "x_0": "x₀",
    "x_i": "xᵢ",
    "y_i": "yᵢ",
    "J(theta)": "J(θ)",
    "J_lambda(theta)": "Jλ(θ)",
    "lambda": "λ",
    "alpha": "α",
    "gamma": "γ",
    "epsilon": "ε",
    "delta": "δ",
    "log2": "log₂",
    "sqrt": "√",
    "sum_i": "∑ᵢ",
    "sum_v": "∑ᵥ",
    "sum p_i": "∑ᵢ pᵢ",
    "sum_i p_i": "∑ᵢ pᵢ",
    "sum p_i^2": "∑ᵢ pᵢ²",
    "sum_i p_i^2": "∑ᵢ pᵢ²",
    "p_i": "pᵢ",
    "p1": "p₁",
    "n1": "n₁",
    "p0": "p₀",
    "n0": "n₀",
    "p+": "p₊",
    "p-": "p₋",
    "S_v": "Sᵥ",
    "z_N": "z_N",
    "error_S": "err_S",
    "error_D": "err_D",
    "Var(epsilon)": "Var(ε)",
    "Bias(f_hat(x0))": "Bias(f̂(x₀))",
    "Var(f_hat(x0))": "Var(f̂(x₀))",
    "f_hat(x0)": "f̂(x₀)",
    "f_hat": "f̂",
    "y0": "y₀",
    "x0": "x₀",
    "A1": "A₁",
    "A2": "A₂",
    "Ak": "Aₖ",
    "not A1": "¬A₁",
    "not A2": "¬A₂",
    "not B": "¬B",
    "->": "→",
    "!=": "≠",
    "<=": "≤",
    ">=": "≥",
}


EXACT_REPLACEMENTS = {
    "Entropy(S)=-p+log2(p+)-p-log2(p-). Entropy bằng 0 khi tập hoàn toàn cùng lớp và cao khi các lớp trộn lẫn.":
        "Entropy(S) = −p₊log₂(p₊) − p₋log₂(p₋). Entropy bằng 0 khi tập hoàn toàn cùng lớp và cao khi các lớp trộn lẫn.",
    "Gain(S,A)=H(S)-sum_v |S_v|/|S| H(S_v)": "Gain(S,A) = H(S) − ∑ᵥ (|Sᵥ|/|S|)H(Sᵥ)",
    "H(S)=-sum p_i log2(p_i)": "H(S) = −∑ᵢ pᵢ log₂(pᵢ)",
    "-sum p_i log2 p_i": "−∑ᵢ pᵢ log₂(pᵢ)",
    "1 - sum p_i^2": "1 − ∑ᵢ pᵢ²",
    "Gini(S)=1-sum_i p_i^2": "Gini(S) = 1 − ∑ᵢ pᵢ²",
    "E[(y0 - f_hat(x0))^2] = Var(f_hat(x0)) + [Bias(f_hat(x0))]^2 + Var(epsilon).":
        "E[(Y₀ − f̂(x₀))²] = Var(f̂(x₀)) + Bias(f̂(x₀))² + Var(ε).",
    "E[(y0-f_hat(x0))^2]=Var(f_hat)+Bias(f_hat)^2+Var(epsilon)":
        "E[(Y₀ − f̂(x₀))²] = Var(f̂(x₀)) + Bias(f̂(x₀))² + Var(ε)",
    "error_S ± z sqrt(error_S(1-error_S)/n)": "err_S(h) ± z_N√(err_S(h)(1−err_S(h))/n)",
    "error_S ± z_N sqrt(error_S(1-error_S)/n)": "err_S(h) ± z_N√(err_S(h)(1−err_S(h))/n)",
    "error_D(h)=Pr_{x~D}[h(x) != c(x)]": "err_D(h) = Prₓ∼D[h(x) ≠ c(x)]",
    "error_S(h)=số mẫu bị phân loại sai / |S|": "err_S(h) = số mẫu bị phân loại sai trên S / |S|",
    "m tăng theo (1/epsilon)(ln|H|+ln(1/delta))": "m = O((1/ε)(ln|H| + ln(1/δ)))",
    "t[log2(p1/(p1+n1))-log2(p0/(p0+n0))]": "t[log₂(p₁/(p₁+n₁)) − log₂(p₀/(p₀+n₀))]",
    "FOIL-Gain(L,R) = t * [log2(p1/(p1+n1)) - log2(p0/(p0+n0))].":
        "FOIL-Gain(L,R) = t[log₂(p₁/(p₁+n₁)) − log₂(p₀/(p₀+n₀))].",
}


PREFIX_REWRITES = [
    (
        "Entropy: Entropy đo độ hỗn loạn của tập dữ liệu. Với hai lớp,",
        "Entropy: Entropy đo độ hỗn loạn của tập dữ liệu. Với hai lớp, công thức là Entropy(S) = −p₊log₂(p₊) − p₋log₂(p₋). Entropy bằng 0 khi tập hoàn toàn cùng lớp và cao khi các lớp trộn lẫn."
    ),
    (
        "Information Gain: Information Gain đo mức giảm entropy sau khi chia theo thuộc tính A:",
        "Information Gain: Information Gain đo mức giảm entropy sau khi chia theo thuộc tính A: Gain(S,A) = H(S) − ∑ᵥ (|Sᵥ|/|S|)H(Sᵥ). Thuộc tính có gain cao thường được chọn trước trong ID3."
    ),
    (
        "Công thức chuẩn: E[",
        "Công thức chuẩn: E[(Y₀ − f̂(x₀))²] = Var(f̂(x₀)) + Bias(f̂(x₀))² + Var(ε)."
    ),
    (
        "Ý nghĩa từng thành phần: Var(",
        "Ý nghĩa từng thành phần: Var(f̂(x₀)) đo độ nhạy của mô hình đối với tập huấn luyện; Bias(f̂(x₀)) đo sai lệch hệ thống giữa dự đoán trung bình của mô hình và hàm thật; Var(ε) là nhiễu không thể giảm, còn gọi là irreducible error."
    ),
    (
        "Công thức xấp xỉ: error_D",
        "Công thức xấp xỉ: err_D(h) nằm trong khoảng err_S(h) ± z_N√(err_S(h)(1−err_S(h))/n). Trong đó n là số mẫu test, z_N phụ thuộc mức tin cậy: khoảng 1.64 cho 90%, 1.96 cho 95%, 2.58 cho 99%."
    ),
    (
        "Ví dụ tính nhanh: Nếu error_S=0.12",
        "Ví dụ tính nhanh: Nếu err_S(h)=0.12, n=100, mức tin cậy 95% dùng z_N=1.96. Độ lệch = 1.96√(0.12×0.88/100) ≈ 1.96×0.0325 ≈ 0.064. Vậy err_D(h) xấp xỉ nằm trong 0.12 ± 0.064, tức từ 0.056 đến 0.184."
    ),
    (
        "Công thức trực giác cho H hữu hạn:",
        "Công thức trực giác cho H hữu hạn: số mẫu m = O((1/ε)(ln|H| + ln(1/δ))). Không gian giả thuyết càng lớn, hoặc yêu cầu ε và δ càng nhỏ, thì cần càng nhiều dữ liệu."
    ),
]


def normalize_font(doc: Document):
    for style_name in ["Normal", "Heading 1", "Heading 2", "Heading 3"]:
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        if style_name == "Normal":
            style.font.size = Pt(12)


def normalize_text(text: str) -> str:
    new = text
    for prefix, replacement in PREFIX_REWRITES:
        if new.startswith(prefix):
            return replacement
    for old, replacement in EXACT_REPLACEMENTS.items():
        new = new.replace(old, replacement)
    for old, replacement in DIRECT_REPLACEMENTS.items():
        new = new.replace(old, replacement)
    # Cleanup artifacts created by chained replacements.
    new = new.replace("hθ(x)", "h_θ(x)")
    new = new.replace("Jλ(θ)", "J_λ(θ)")
    new = new.replace("Prₓ∼D", "Pr_{x∼D}")
    new = new.replace("z_N", "z_N")
    new = new.replace("1−err_S", "1 − err_S")
    new = new.replace("(1-err_S", "(1 − err_S")
    new = new.replace(")H(Sᵥ)", ") H(Sᵥ)")
    new = new.replace("−∑", "−∑")
    return new


def set_paragraph_text(paragraph, text):
    if paragraph.text == text:
        return False
    if not paragraph.runs:
        paragraph.add_run(text)
        return True
    for run in paragraph.runs:
        run.text = ""
    paragraph.runs[0].text = text
    return True


def normalize_doc(doc: Document) -> int:
    changed = 0
    for para in doc.paragraphs:
        text = para.text
        new = normalize_text(text)
        if new != text:
            changed += set_paragraph_text(para, new)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    text = para.text
                    new = normalize_text(text)
                    if new != text:
                        changed += set_paragraph_text(para, new)
    return changed


def main():
    if not TARGET.exists():
        raise FileNotFoundError(TARGET)
    if not BACKUP.exists():
        copy2(TARGET, BACKUP)

    doc = Document(TARGET)
    normalize_font(doc)
    changed = normalize_doc(doc)

    try:
        doc.save(TARGET)
        print(f"UPDATED={TARGET}")
    except PermissionError:
        doc.save(OUTPUT_IF_LOCKED)
        print(f"TARGET_LOCKED={TARGET}")
        print(f"OUTPUT={OUTPUT_IF_LOCKED}")
    print(f"BACKUP={BACKUP}")
    print(f"CHANGED_PARAGRAPHS_OR_CELLS={changed}")


if __name__ == "__main__":
    main()
